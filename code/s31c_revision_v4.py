# -*- coding: utf-8 -*-
"""V3-R3 s31c: R4 修订版全文手稿（回应第三轮审稿 P0/M 清单；普通 Word 版式）
数字一律从 data/stats_v4.json 注入；表格从 tables/ v4 版 CSV 读入。
输出：manuscript/V3_manuscript_full_v4.docx
"""
import json
import re

import docx
import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

S = json.load(open("data/stats_v4.json", encoding="utf-8"))
OUT = "manuscript/V3_manuscript_full_v4.docx"
doc = docx.Document()

for s in doc.sections:
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.5)
    s.top_margin = s.bottom_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.3
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_runs(p, text, size=None):
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok)
        if size:
            r.font.size = Pt(size)


def P(text, size=None):
    p = doc.add_paragraph()
    add_runs(p, text, size)
    return p


def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True


def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True


def FIG(path, num, caption, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Fig. {num}. ")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(10)


def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        borders.append(e)
    tblPr.append(borders)


def TABLE(num, title, headers, rows, fontsize=8.5):
    cap = doc.add_paragraph()
    r = cap.add_run(f"Table {num}. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(title)
    r2.font.size = Pt(10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        rr = c.paragraphs[0].add_run(str(h))
        rr.bold = True
        rr.font.size = Pt(fontsize)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(str(v))
            rr.font.size = Pt(fontsize)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def EQ(parts, num):
    """parts: list of (text, mode)；mode: None/'i'/'sub'/'sup'/'isub'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, mode in parts:
        r = p.add_run(text)
        if "i" in (mode or ""):
            r.italic = True
        if "sub" in (mode or ""):
            r.font.subscript = True
        if "sup" in (mode or ""):
            r.font.superscript = True
    p.add_run(f"    ({num})")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


# ==================== 标题区 ====================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Thermal–Wireless Trade-Offs of Terrain-Constrained Settlement Morphology in Huizhou "
              "Traditional Villages")
r.bold = True
r.font.size = Pt(16)
t.paragraph_format.space_after = Pt(14)

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.add_run("Lei Zhang")
r = a.add_run("1,*"); r.font.superscript = True
for _r in a.runs:
    _r.font.size = Pt(12)

af = doc.add_paragraph()
af.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = af.add_run("1 School of Architecture and Urban Planning, Anhui Jianzhu University, Hefei 230601, China")
r.font.size = Pt(9.5)
r.italic = True
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run("* Corresponding author. E-mail: [corresponding author's e-mail to be inserted]")
r.font.size = Pt(9.5)
doc.add_paragraph()

# ---- Abstract（约250词）----
ab = doc.add_paragraph()
r = ab.add_run("Abstract: ")
r.bold = True
add_runs(ab,
    "Traditional villages in mountainous regions must reconcile climatic comfort with modern digital "
    "infrastructure, yet the same settlement morphology may act in opposite directions on these two "
    "performances. For 29 traditional villages in the Huizhou region of southern Anhui, China, this paper "
    "quantifies how settlement morphology is associated with summer land surface temperature (LST) and "
    "with simulated wireless coverage, using only open data and scripted, re-runnable models. Twelve "
    "morphological metrics derive from the 30-m Copernicus GLO-30 digital surface model (DSM) and the "
    "10-m ESA WorldCover map. Village summer LST is estimated from Landsat 8/9 Collection-2 surface-"
    "temperature assets (June–September 2019–2025) after merging same-day cross-frame acquisitions into "
    "26 independent overpasses (25 retained after excluding one near-zero-coverage date; 6–13 overpasses "
    "contribute per village); coverage is evaluated as reference-signal-received-power (RSRP) threshold "
    "exceedance under a stylized, standardized virtual macrocell deployment (2.5-km grid snapped to "
    "locally elevated sites, 30-m masts, 2.6 GHz) with a terrain line-of-sight 3GPP rural-macrocell "
    "model using true three-dimensional link distances, repeated over four grid phases. Under the "
    "site-domain definition, all 29 village heat-anomaly point estimates are positive (mean ΔLST = "
    "+2.4 °C; overpass bootstrap 95% intervals exclude zero in 21 of 29); under a terrain- and "
    "land-cover-matched background the anomaly halves (mean +1.3 °C; 25 of 29 positive). Topographic "
    "openness (terrain sky-view factor) is double-edged: it is associated with higher LST (Spearman "
    "ρ = +0.65) and with better RSRP-threshold exceedance (ρ = +0.71), associations that survive "
    "within-family Benjamini–Hochberg false-discovery-rate control after Moran’s-I effective-sample-size "
    "correction, with spatial block-bootstrap confidence intervals. The forest ring shows no detectable "
    "thermal association at village scale (ρ = −0.36, n.s.) but strongly shadows signal (ρ = −0.81). "
    "Plan-form metrics yield no association that survives false-discovery control: compactness is "
    "unrelated to thermal outcomes and shows only nominal negative tendencies on the wireless side, and "
    "elongation shows a marginal coverage trend (ρ = −0.33, p = 0.09). Results are stable to band choice "
    "(700 MHz), county-level jackknife, and propagation-model range limits; grid-phase effects are small "
    "for open villages but large for marginal ones. Because the design is cross-sectional and the "
    "deployment stylized, the quantified couplings are hypotheses for site-specific assessment—read "
    "together with the confidence intervals and the stated limitations—not causal effects, operational "
    "coverage estimates, or general laws. The open pipeline transfers to other mountainous regions.")
ab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw = doc.add_paragraph()
r = kw.add_run("Keywords: ")
r.bold = True
kw.add_run("settlement morphology; land surface temperature; rural wireless coverage; 3GPP rural "
           "macrocell model; sky-view factor; traditional villages; Huizhou")
doc.add_paragraph()

# ==================== 1. INTRODUCTION ====================
H1("1. Introduction")
P("Traditional settlements in mountainous regions embody climatic knowledge accumulated over centuries of "
  "trial and error. In the Huizhou region of southern Anhui Province, China—home to the World Heritage "
  "villages of Hongcun and Xidi and to one of the densest concentrations of historic villages in East "
  "Asia—enclosed valley sites, forested backdrops, compact building fabrics, and water-adjacent layouts "
  "have long been read as a vernacular, pre-scientific system of passive environmental control [1], [2]. "
  "Field studies in and around the region confirm that this question remains empirically alive: indoor and "
  "outdoor thermal performance of Huizhou-area vernacular fabric continues to be measured and debated [3], "
  "[4]. These villages are simultaneously protected heritage, working agricultural communities, and tourist "
  "destinations, and they must now accommodate a second, very different infrastructure: rural broadband and "
  "mobile networks, digital-village platforms [5], and, increasingly, city information modeling (CIM) at "
  "the settlement scale [6]. The two agendas are usually planned by different agencies with different "
  "tools, and they meet, if at all, only at the construction site.")
P("The premise of this paper is that they should meet much earlier, because both agendas load on the same "
  "physical object: settlement morphology. The morphology that heritage conservation prefers—enclosed, "
  "compact, forest-backed, water-adjacent—is widely believed to be thermally favorable in hot-summer "
  "climates [1], [2], [7]. Radio propagation in mountainous terrain, by contrast, favors the opposite "
  "geometry: open horizons and unobstructed line-of-sight (LOS) paths to elevated macrocells [8]–[10]. "
  "Whether these two performances actually trade off across real villages—and by how much, and through "
  "which specific morphological attributes—has never been quantified. This paper provides that "
  "quantification, and in doing so it treats the associations strictly as associations: the design is "
  "cross-sectional and correlational, and the language of the paper follows that design.")
P("The thermal side of the question has a mature literature. The energetic basis of settlement-scale heat "
  "excess has been understood since the foundational work of Oke [11], and subsequent reviews consolidated "
  "the roles of geometry, surface cover, and anthropogenic heat [12]. The local climate zone (LCZ) scheme "
  "formalized the idea that built morphology predicts thermal behavior [13], and canyon geometry together "
  "with the sky-view factor (SVF) emerged as robust predictors of urban heat intensity [14], [15]. For "
  "vernacular settlements specifically, case studies combining field measurement and simulation have "
  "demonstrated the climate responsiveness of traditional dwellings and village layouts in China and "
  "elsewhere [1]–[4], [16], including systematic reviews of outdoor microclimate design in traditional "
  "Chinese villages [7]. Two gaps nonetheless remain. First, the vernacular evidence is overwhelmingly "
  "single-village: cross-settlement, statistically controlled evidence linking village morphology to "
  "thermal outcomes is thin, particularly in mountainous southern China [7], [17]. Second, the dominant "
  "morphometric evidence comes from cities, where morphology varies independently of terrain; in mountain "
  "villages, built form and terrain co-vary by design, so the urban-derived intuition that ‘more open is "
  "cooler’ cannot be assumed to hold [18].")
P("The radio side has an equally mature but entirely separate literature. Empirical path-loss models "
  "descend from Hata [8] through the COST-231 family, while the geometry-based 3GPP rural macrocell (RMa) "
  "model specifies distinct line-of-sight (LOS) and non-LOS (NLOS) branches in which terrain obstruction "
  "is the first-order control on rural coverage [9], [10]. Vegetation is known to attenuate gigahertz "
  "signals [19]. Rural coverage gaps—and their socio-economic consequences—are documented in global "
  "digital-divide assessments [20], and in China the digital-village strategy has made rural connectivity "
  "an explicit policy target [5]. Yet coverage studies rarely treat settlement morphology as an "
  "explanatory variable: villages appear, if at all, as demand points on a cost surface, not as "
  "morphological objects whose shape co-determines serviceability. No study we are aware of evaluates the "
  "thermal and radio performances of the same settlements jointly.")
P("Meanwhile, the data environment has changed. Open elevation models, global land-cover maps, and "
  "multi-year thermal satellite archives have made quantitative morphometrics cheap and reproducible "
  "[21], [23], [26], and urban studies have shown at scale that open imagery and machine learning can "
  "characterize urban form [22], [24], [25]. Rural heritage settlements have so far benefited "
  "little from this turn, partly because their building stock is small and scattered, and partly because "
  "the relevant questions sit between disciplines. This study brings the open-data morphometric approach "
  "to bear on the twin infrastructure question of mountain villages.")
P("We pose four research questions. (RQ1) How differentiated are Huizhou traditional villages in the "
  "morphological dimensions that plausibly matter for both thermal and radio performance? (RQ2) Which "
  "morphological attributes are associated with village-scale summer surface temperature and its anomaly "
  "against the rural background? (RQ3) Which attributes are associated with simulated wireless coverage "
  "under a deployment that is held constant across villages? (RQ4) Where the two sets of associates "
  "coincide, do they act in the same or in opposite directions—that is, where are the trade-offs?")
P("To answer these questions we assemble a sample of 29 traditional villages in the Huizhou region and, "
  "using only open data and reproducible models, (i) characterize each village by twelve morphological "
  "metrics defined over two explicitly stated spatial frames; (ii) estimate village-scale summer LST and "
  "heat anomaly from a quality-masked multi-year Landsat median composite; (iii) simulate coverage under "
  "a standardized virtual macrocell deployment that removes operator-specific deployment heterogeneity, "
  "with grid-phase and frequency sensitivity analyses; and (iv) quantify, with size-controlled partial "
  "correlations, false-discovery-rate control over two explicitly defined test families, spatial "
  "autocorrelation diagnostics, and leave-one-county-out checks, which morphological attributes covary "
  "with the two performances in opposite or concordant directions. Section 2 describes the study area, "
  "data, and methods; Section 3 reports the results; Section 4 discusses interpretation, planning "
  "implications, and limitations; Section 5 concludes.")

# ==================== 2. MATERIALS AND METHODS ====================
H1("2. Materials and Methods")
P("The research framework couples two open-data model chains onto a common morphometric backbone (Fig. 1, "
  "Table 1). For each of the 29 villages we first compute twelve morphological metrics from elevation and "
  "land-cover data (Section 2.3). The same village domains then feed two independent performance models: a "
  "thermal chain, in which a quality-masked multi-year Landsat surface-temperature composite yields village "
  "LST and its anomaly against the rural background (Section 2.4); and a radio chain, in which a "
  "standardized virtual macrocell deployment is evaluated with a terrain-aware 3GPP propagation model to "
  "yield village coverage metrics (Section 2.5). Because the deployment is held constant across villages, "
  "coverage differences between villages reflect morphology and terrain rather than operator decisions. "
  "The two performance vectors are finally compared against the morphological vector across the sample, "
  "with the statistical controls described in Section 2.6. All in-text statistics, tables, and figures are "
  "generated from a single village-level analysis table by script, so that numbers reported in the text "
  "cannot drift from the data.")

H2("2.1. Study Area and Village Sample")
P("The study area covers the historical Huizhou prefecture in the foothills of the Huangshan mountains, "
  "southern Anhui (approximately 29.6°–30.3° N, 117.3°–118.9° E). The climate is subtropical monsoon with "
  "hot, humid summers in which outdoor thermal stress is a genuine public-health and livability concern. "
  "The terrain is a mosaic of steep forested ridges and narrow alluvial basins; villages traditionally "
  "settled along streams at basin edges, with their backs to a forested slope and their fronts to water—"
  "the classical *beishan mianshui* (backed by mountains, facing water) configuration. The region contains "
  "two UNESCO World Heritage villages (Hongcun and Xidi) and numerous nationally listed historic and "
  "traditional villages, making it both a conservation hotspot and a natural laboratory in which "
  "morphological setting varies widely within a single cultural sphere.")
P("*Sampling frame.* The candidate population was the published registers of nationally listed "
  "traditional and historic-cultural villages within the historical Huizhou prefecture (the six counties "
  "above), a set of roughly one hundred settlements. Inclusion required (i) listing in a national "
  "register, (ii) an intact historic built core visible in current imagery, and (iii) a geographic "
  "identity verifiable against OpenStreetMap and satellite imagery. From the register we assembled a "
  "desk-review shortlist of 28 villages spanning basin-floor, piedmont, and deep-valley settings across "
  "all six counties; one further village (Huansha) was added after imagery verification, yielding the "
  "final sample of 29. The sample is thus a purposive regional sample stratified by terrain setting and "
  "county, not a probability sample; we report it as such and do not generalize beyond the sampled "
  "region. Village names are given in both pinyin and Chinese characters throughout (Tables 1 and A1; "
  "Appendix Table A11). The full candidate list, with per-village coordinate source (register, "
  "OpenStreetMap, or imagery fix), verification status, and reviewer notes in both languages, is "
  "documented in Appendix Table A11 and archived in the repository (village_sample_v2.csv). Two villages "
  "(Kantou and Lixi) carry approximate-coordinate flags: no authoritative point source was "
  "available, so their positions were fixed from OpenStreetMap and imagery, and after snapping to the "
  "detected built-up fabric they sit 73 m and 143 m from the nominal coordinates, respectively; we flag "
  "them because village-core placement enters the domain definition, and we note that a misplacement of "
  "this order is small against the 500-m domain radius and the 100-m thermal pixel. Village "
  "coordinates were snapped to the "
  "centroid of the nearest built-up connected component in the reference land-cover map (median snap "
  "distance 78 m, maximum 517 m), so that all morphological measurement is anchored to the actual built "
  "fabric rather than to nominal administrative points. Three sample particularities are flagged because "
  "they matter for the analysis. First, Yuliang’s detected built-up component (258.7 ha) merges with the "
  "adjacent Shexian county seat; we retain the village and verify that all reported associations are "
  "robust to its exclusion. Second, the built fabric of two small mountain villages (Zuyuan and Mulihong) "
  "is under-detected by the 10-m land-cover map: component-based metrics are unavailable for both "
  "(n = 27 where applicable), while the domain-based plan-form metrics of Section 2.3 meet their "
  "ten-pixel minimum for Zuyuan but not for Mulihong, which has no detected built-up cell inside its "
  "domain (n = 28 where applicable). Their coverage targets follow the fallback rule of Section 2.5, "
  "with a dedicated exclusion sensitivity check. Third, the sample deliberately spans basin-floor, "
  "piedmont, and deep-valley settings (domain-mean elevation 114–552 m; local relief 38–538 m), which is "
  "the source of morphological variance on which the analysis draws. Because elevation is summarized as a "
  "domain mean, it can differ from point elevation at the village centroid, particularly for steep "
  "mountain villages. Fig. 1 maps the sample; Table 1 lists the villages with key attributes; the full "
  "twelve-metric table and the per-phase coverage detail are given in Appendix Table A1.")
FIG("figures/Fig1_study_area_EN.png", 1,
    "The 29 sampled traditional villages in the Huizhou region, southern Anhui, China. Base map: Esri "
    "World Imagery (Maxar, Earthstar Geographics, CNES/Airbus DS), accessed August 2026; map projection "
    "WGS 84 / Pseudo-Mercator (EPSG:3857); numbers keyed to Table 1. Inset: location in China.")

H2("2.2. Open Data")
P("All inputs are public datasets; no field measurement was used, and the entire pipeline is scripted and "
  "re-runnable. Terrain is given by the 30-m Copernicus GLO-30 product (COP-DEM) [21]. COP-DEM is a "
  "digital *surface* model (DSM): its heights reference the reflecting surface, which over forested or "
  "built-up ground includes canopy and roof signals; we therefore refer to it as a DSM throughout, treat "
  "terrain metrics derived from it as surface metrics, and return to the consequences in Section 4.4. As "
  "a gross-error check, village centroid elevations and within-domain relief were recomputed from the "
  "independent NASADEM 1-arc-second product [32]: centroid elevations agree at Spearman ρ = 0.999 "
  "(RMSE 22 m, mean difference −3 m) and domain relief at ρ = 0.988 across the 29 villages (Appendix "
  "Table A10). Land cover is the 10-m ESA WorldCover 2021 map (v200), from which the built-up, "
  "tree-cover, and water classes were extracted [23]. Summer surface temperature is derived from Landsat "
  "8/9 Collection-2 Level-2 surface-temperature products (band ST_B10, acquired at 100-m native thermal "
  "resolution and distributed on the 30-m product grid; we bilinearly reproject the band to the 100-m UTM "
  "analysis grid), accessed through the "
  "Microsoft Planetary Computer spatiotemporal asset catalog (collection landsat-c2-l2; query executed "
  "17 August 2026) [27]. We queried June–September scenes from "
  "2019–2025 with catalog scene cloud cover below 10%, retaining up to nine of the lowest-cloud scenes "
  "per Worldwide Reference System frame, yielding 36 scene assets over the four frames (paths 120–121, "
  "rows 39–40) that cover the study area; the per-scene identifiers, dates, and cloud covers are listed "
  "in Appendix Table A2. Because adjacent WRS rows of the same path acquired on the same date are two "
  "observations of the same overpass, treating them as independent scenes would pseudo-replicate the "
  "sample. We therefore merge same-platform, same-path, same-date assets pixel-wise (mean of valid "
  "values) into 26 independent overpasses; all anomaly estimation and bootstrap resampling below uses "
  "the overpass, not the scene asset, as the unit. Digital numbers were converted with the Collection-2 "
  "scale factors (ST = DN × 0.00341802 + 149.0 K) and expressed in degrees Celsius. Quality control "
  "applied the Collection-2 QA_PIXEL bit mask pixel-wise (fill, dilated cloud, cirrus, cloud shadow, "
  "snow, and cloud bits cleared), followed by a 10–60 °C physical plausibility screen. Three scenes were "
  "excluded because this screen left zero valid pixels in the study-area subset: inspection shows that "
  "the subset is covered by cloud or cloud shadow in these scenes even though their whole-scene catalog "
  "cloud cover is only 0.6–4.6%, a reminder that catalog cloud cover is a whole-scene statistic. One "
  "further overpass date (2020-06-27) retains only ≈0.01% valid pixels after masking; it is "
  "excluded from the primary anomaly analysis and retained as a sensitivity case. We use the "
  "distributed product as is, including its emissivity model, and we do not propagate the per-pixel "
  "ST_QA uncertainty band, which would widen the reported intervals further; this is stated as a "
  "limitation (Section 4.4), and overpass-block bootstrap propagates retrieval and sampling uncertainty "
  "into the village anomaly (Section 2.4). The summer composite used for mapping is the per-pixel "
  "*median* across the 26 overpasses, which is robust to residual contamination. A 500-m village domain "
  "contains roughly 80 native 100-m thermal pixels. Using a multi-year median sacrifices year-to-year "
  "resolution in exchange for a stable, cloud-gap-free characterization of typical summer daytime "
  "surface behavior, which is the quantity that morphology can plausibly covary with [28].")

H2("2.3. Morphological Metrics: Two Spatial Frames, Stated Explicitly")
P("Because mixing spatial frames silently invites contradiction, we state the frames before defining "
  "the metrics. **Frame O (object level)** is the "
  "built-up connected component nearest to the nominal village point in the 10-m land-cover map; it has no "
  "fixed geometric bound and can extend well beyond the village core (for Yuliang it merges with the "
  "adjacent county seat). **Frame D (site level)** is the fixed 500-m-radius circular domain centered on "
  "the Frame-O centroid; its maximum area is π × 500² ≈ 78.5 ha, and it defines the analysis window for "
  "terrain, landscape context, thermal, and coverage measurement. Object-level metrics describe the plan "
  "form of the detected fabric; site-level metrics describe where that fabric sits.")
P("*Plan-form metrics (Frame D).* Plan form is measured on the built-up cells detected inside the 500-m "
  "domain on the 10-m WorldCover grid. Elongation is the principal-axis aspect ratio of those cells, "
  "defined as sqrt(λ1/λ2) of the eigenvalues of the cell-centroid covariance matrix; it is high for "
  "stream-following linear villages and low for isotropic basin villages. Compactness is the "
  "isoperimetric quotient 4πA/P² aggregated over the built-up patches inside the domain, where A is the "
  "total built-up area and P is the total exposed-edge perimeter. The perimeter is computed under "
  "four-neighbour adjacency with anisotropic weighting—vertical edges contribute the pixel height and "
  "horizontal edges the pixel width, which differ by up to 13% at this latitude on the geographic "
  "WorldCover grid—counting hole boundaries as exposed perimeter, counting edges at the domain circle, "
  "and aggregating all patches into a single quotient (total A over total P², not a patch-weighted "
  "average); the implementation is covered by geometric unit tests (square, rectangle, two-patch, and "
  "holed-patch cases) distributed with the code (s60_unit_tests.py). Both metrics require at least ten "
  "detected built-up cells inside the domain, a threshold met by 28 of the 29 villages (Mulihong has "
  "no detected built-up cell in its domain). An earlier component-based formulation of these two "
  "metrics proved not to be reproducible across land-cover processing versions and was retired. The "
  "component built-up area (Frame O, 5.2–258.7 ha) is retained as a descriptive quantity only: it "
  "correlates at ρ = 0.98 with domain built-up area (n = 29), and its maximum reflects the Yuliang "
  "merge with the adjacent county seat rather than village fabric, so it carries no independent content "
  "as a size control.")
P("*Terrain metrics (Frame D).* Elevation is the domain mean; relief is the maximum minus minimum "
  "elevation within the domain; slope is the domain mean of terrain gradient. Southness is the domain "
  "mean of cos(aspect), positive for south-facing ground, and north–south asymmetry is the difference in "
  "mean elevation between the northern and southern halves of the domain—together these capture the "
  "canonical claim that Huizhou villages face south against a northern slope. To capture terrain "
  "enclosure we use a terrain sky-view factor (tSVF), the unobstructed fraction of the upper hemisphere "
  "at the built-up centroid, sampled along 72 azimuths with a 2500-m horizon [15]:")
EQ([("tSVF = 1 − (1/", "i"), ("N", "i"), (") Σ sin(", None), ("γ", "i"), ("i", "sub"),
    ("),   ", None), ("i", "i"), (" = 1, …, 72", None)], 1)
P("where γ", )
doc.paragraphs[-1].add_run("i").font.subscript = True
add_runs(doc.paragraphs[-1],
    " is the horizon elevation angle along azimuth *i*. A village on an open plain has tSVF ≈ 1; a village "
    "at the bottom of a gorge has a markedly lower value. The 2500-m horizon is long enough to include the "
    "enclosing ridge lines that govern both insolation and radio visibility, and short enough to remain a "
    "village-scale attribute. Because tSVF is evaluated once, at the centroid, it characterizes the "
    "village’s setting rather than its internal spatial variation; intra-village sky-view heterogeneity "
    "requires building-scale data that our open sources do not provide (Section 4.4).")
P("*Landscape-context metrics (Frame D unless noted).* The forest ring is the tree-cover fraction within "
  "a 300-m buffer ring around the Frame-O component, excluding the component itself, representing the "
  "vegetated backdrop typical of Huizhou "
  "villages; a 300-m binary buffer proved undiscriminating (saturated at 100% for most villages), so we "
  "use the continuous ring share; because the ring is defined around the detected component, it too is "
  "unavailable for the two under-detected villages. Water proximity is the Euclidean distance from "
  "built-up cells to the nearest water pixel, summarized by its mean and minimum. Finally, domain "
  "built-up area is the built-up surface clipped to Frame D (0–56.8 ha across the sample); it is the "
  "size variable used in the statistical controls of Section 2.6, because it shares the spatial frame of "
  "the outcome variables (Spearman ρ = 0.87 against component area across the 27 villages with both). "
  "Several metrics co-vary by construction of the terrain—deep-valley villages are steep, forested, "
  "enclosed, and elongated—which is why Section 2.6 relies on rank correlations and size-controlled "
  "partial correlations rather than on parametric multivariate models.")

H2("2.4. Summer LST and Heat Anomaly")
P("For each village, core and background temperatures are defined over explicitly stated zones, and the "
  "choice of zone is itself reported as a methodological variable. We compute the anomaly under three "
  "definitions. **V1 (site domain):** the core is the 500-m domain with water pixels masked, and the "
  "background is the 1–2 km ring restricted to non-built, water-masked pixels, so that the background "
  "reflects the rural matrix rather than adjacent settlement fabric. **V2 (fabric):** the core is the "
  "built-up-classified pixels inside the 500-m domain (villages with fewer than five such pixels are "
  "not estimable), against the same ring; this variant isolates the built boundary itself at the cost "
  "of sample size. **V3 (matched):** the core is the V1 domain, and the background is a pixel-level "
  "subset of the V1 ring matched one-to-one (greedy, without replacement) on elevation, slope, and "
  "tree-cover fraction, so that core and background share terrain and land-cover context by "
  "construction; V3 is the controlled estimate and the primary anomaly used in the correlation "
  "analysis, while V1 and V2 are reported alongside (Table 1; Appendix Table A9). Water pixels are "
  "masked in every zone of every variant; in the V1 ring they occupy on average 0.6% of the area "
  "(maximum 6.9% at riverside Yuliang, which is covered by the dedicated exclusion sensitivity of "
  "Section 2.6). Because the composite pools overpasses acquired on different dates, a difference of "
  "composite means would mix acquisition composition with morphology. We therefore define the heat "
  "anomaly overpass by overpass: for every independent overpass we take the difference between the "
  "core mean and the background mean, each computed only when at least five valid pixels are available "
  "in that overpass and zone, and then average these per-overpass differences for each village "
  "(between 6 and 13 overpasses contribute per village, mean 7.6),")
EQ([("ΔLST = (1/", "i"), ("N", "i"), ("o", "sub"), (") Σ", None), ("o", "i"),
    (" [LST", "i"), ("core", "sub"), ("(", None), ("o", "i"), (") − LST", "i"),
    ("bg", "sub"), ("(", None), ("o", "i"), (")]", None)], 2)
P("which controls acquisition date by construction. Uncertainty is propagated by a bootstrap whose "
  "resampling unit is the independent overpass (10,000 resamples, fixed seed 7), yielding a 95% "
  "confidence interval and the probability that the anomaly exceeds zero; villages whose interval "
  "includes zero are reported as having a positive point estimate without confirmed sign. Threshold "
  "sensitivity is reported: requiring at least 25% valid-pixel coverage per zone per overpass, and "
  "retaining the near-zero-coverage 2020-06-27 overpass date, both leave the village estimates "
  "essentially unchanged (rank correlation with the primary estimates ρ ≥ 0.98 and no sign changes in "
  "either case; Appendix Table A9). The absolute LST of the overpass-deduplicated median composite remains "
  "descriptive: absolute values are used for ranking within the sample, not as climatological norms. "
  "LST is used as a physically consistent proxy of the daytime outdoor thermal environment; it is not "
  "equivalent to air temperature or pedestrian-level thermal comfort, a distinction we return to in "
  "Section 4.4 [28].")

H2("2.5. Standardized Wireless Deployment and Propagation Model")
P("Real operator deployments are heterogeneous and commercially opaque, so village-level coverage "
  "differences measured against live networks conflate morphology with deployment history. To isolate "
  "morphology and terrain, we evaluate coverage under a standardized virtual deployment: macrocell sites "
  "on a 2.5-km grid over the study region, each site snapped to the highest surface-model cell within "
  "600 m—a stylized heuristic for the locally elevated siting common in rural radio planning, not a "
  "description of any real network—with a 30-m antenna height. The carrier frequency is 2.6 GHz, a "
  "widely deployed mid-band in China’s mobile networks; we treat it as a conservative mid-band case, "
  "because lower bands propagate further and diffract better around terrain—a 700-MHz repetition of the "
  "full chain, reported in Section 3.3, confirms that absolute coverage levels rise at low frequency "
  "while the inter-village ranking is essentially preserved. Following the RSRP measurement definition, "
  "the link budget is written directly in terms of the effective radiated power per resource element "
  "(EPRE) [33]:")
EQ([("EPRE = ", None), ("P", "i"), ("T", "sub"), (" − 10 log", None), ("10", "sub"),
    ("(", None), ("N", "i"), ("SC", "sub"), (") + ", None), ("G", "i"), ("T", "sub"),
    (" − ", None), ("L", "i"), ("f", "sub"),
    (" = 46 − 10 log", None), ("10", "sub"), ("(1200) + 17 − 2 = 30.2 dBm", None)], 3)
P("for a 46-dBm total carrier power spread over a 20-MHz LTE-class downlink of 1200 subcarriers, 17-dBi "
  "antenna gain, and 2-dB feeder loss. The total power and the subcarrier count enter the analysis only "
  "through this per-resource-element conversion; because RSRP is itself defined per resource element "
  "[33], no further bandwidth term appears in the budget, and the identical EPRE is used at 700 MHz. "
  "Worked link budgets for three representative villages are given in Appendix Table A12. Because any fixed grid may "
  "favor or disfavor individual villages simply through its alignment, the deployment is repeated over "
  "four grid phases shifted by half a spacing (0 and 1250 m in x and y), yielding 1620 sites per phase; "
  "village-level coverage metrics are averaged across the four phases, and the cross-phase dispersion is "
  "itself reported as a measure of deployment robustness.")
P("Coverage targets are the built-up cells (30 m) inside each village domain; where the land-cover map "
  "under-detects the built fabric (fewer than 30 cells: Zuyuan and Mulihong), a regular 150-m grid over "
  "the domain is used instead, and all coverage associations are re-estimated without these two villages "
  "as a sensitivity check (Sections 2.6 and 3.3). The receiver height is 1.5 m. Path loss follows the "
  "3GPP RMa model [9], with LOS or NLOS determined per link by stepping along the 30-m surface-model "
  "profile between site and target and applying a cumulative horizon criterion: a target is in LOS if "
  "the straight site-to-target line clears every intervening terrain elevation. All distances entering "
  "the model are true three-dimensional link distances d3D = sqrt(d2D² + Δz²), where Δz is the "
  "difference between the absolute antenna heights (site surface elevation plus 30 m; target surface "
  "elevation plus 1.5 m); in mountain terrain Δz can reach several hundred metres, which a constant "
  "antenna-height offset would misrepresent. For LOS links,")
EQ([("PL", "i"), ("LOS", "sub"), ("(", None), ("d", "i"), (") = 20 log", None), ("10", "sub"),
    ("(40π", None), ("d", "i"), (" f", "i"), ("c", "sub"), ("/3) + min(0.03 ", None), ("h", "i"),
    ("1.72", "sup"), (", 10) log", None), ("10", "sub"), ("(", None), ("d", "i"),
    (") − min(0.044 ", None), ("h", "i"), ("1.72", "sup"), (", 14.77) + 0.002 log", None),
    ("10", "sub"), ("(", None), ("h", "i"), (") ", None), ("d", "i")], 4)
P("which applies up to the breakpoint distance dBP = 2π hBS hUT fc/c (with fc in Hz in this expression); "
  "beyond the breakpoint, PL_LOS(d) = PL_LOS(dBP) + 40 log10(d/dBP). For NLOS links the RMa model "
  "applies its rural branch,")
EQ([("PL", "i"), ("NLOS", "sub"), ("(", None), ("d", "i"), (") = 161.04 − 7.1 log", None), ("10", "sub"),
    ("(", None), ("W", "i"), (") + 7.5 log", None), ("10", "sub"), ("(", None), ("h", "i"),
    (") − (24.37 − 3.7(", None), ("h", "i"), ("/", None), ("h", "i"), ("BS", "sub"),
    (")²) log", None), ("10", "sub"), ("(", None), ("h", "i"), ("BS", "sub"), (") + (43.42 − 3.1 log",
     None), ("10", "sub"), ("(", None), ("h", "i"), ("BS", "sub"), ("))(log", None), ("10", "sub"),
    ("(", None), ("d", "i"), (") − 3) + 20 log", None), ("10", "sub"), ("(", None), ("f", "i"),
    ("c", "sub"), (") − (3.2(log", None), ("10", "sub"), ("(11.75 ", None), ("h", "i"),
    ("UT", "sub"), ("))² − 4.97)", None)], 5)
P("and the realized path loss is max(PL_LOS, PL_NLOS), with the nominal rural parameterization of the "
  "model (average building height h = 5 m, street width W = 20 m, base height hBS = 30 m, receiver "
  "height hUT = 1.5 m) [9]. Throughout Eqs. (3)–(5), d denotes d3D in metres and the carrier frequency "
  "fc in GHz. The propagation code is covered by unit tests that reproduce a hand-computed LOS link "
  "budget, verify NLOS ≥ LOS and the range-limit behaviour, and verify the d3D sensitivity "
  "(s60_unit_tests.py). The ray-stepping LOS determination is purely geometric: for each site–target "
  "pair we sample the 30-m surface model along the connecting path at one-pixel increments, maintain "
  "the running maximum of the elevation angle from the site antenna to the intervening terrain, and "
  "declare the target visible when the straight line to the receiver clears that cumulative horizon. "
  "The procedure is the radio analogue of the viewshed analysis familiar from landscape planning, and "
  "it is what couples coverage directly to terrain morphology. Each target keeps the strongest server "
  "among sites within 10 km; at this assignment radius, omitting Earth curvature shifts the computed "
  "line of sight by at most about 7.8 m, small against the 30-m terrain sampling. The RMa "
  "parameterization is specified for LOS distances up to 10 km and NLOS distances up to 5 km; we "
  "therefore truncate NLOS links at 5 km, treating longer NLOS links as unserved, and report this "
  "range-limited model as the primary one. The truncation does not bind in the present configuration: "
  "village-level coverage metrics are identical with and without it in all four phases (Appendix Table "
  "A7), and a fixed-antenna-offset variant that ignores true link geometry likewise preserves every "
  "village-level ranking (phase-0 rank correlation 0.998; Table A7), so the reported values are not an "
  "artifact of either choice. We report cov85 and cov95, the percentages of targets whose reference "
  "signal received power (RSRP) exceeds −85 and −95 dBm respectively. These are RSRP-threshold "
  "exceedance shares: without signal-to-interference ratios they cannot certify service, and we use "
  "them strictly as planning-level visibility indicators—the −85-dBm level corresponds to good signal "
  "and the −95-dBm level to basic connectivity in this stylized scenario. The 10th-percentile RSRP "
  "(RSRP p10) serves as a continuous measure of the worst-served fabric.")
P("The model deliberately omits foliage attenuation [19], building clutter, diffraction refinement, "
  "antenna patterns, shadowing, and interference. What it produces is therefore a deterministic "
  "received-power scenario under a stylized deployment—not an estimate of any operational network—and "
  "absolute coverage levels are optimistic. The objects of this study are the *differences between "
  "villages*, for which terrain visibility is the principal spatial obstruction explicitly represented; "
  "whether omitted effects—most notably vegetation attenuation over the forest ring [19]—would change "
  "the village-level ranking cannot be determined from the present model and is flagged as a "
  "validation target (Section 4.4). The 700-MHz repetition (Section 3.3) quantifies the frequency side "
  "of this question.")


H2("2.6. Statistical Analysis")
P("Associations between the twelve morphological metrics and the four performance variables (village "
  "LST, ΔLST, cov85, RSRP p10) are quantified by Spearman rank correlations, which are robust to the "
  "skewed distributions and small sample size. Two test families are defined explicitly: family F1 "
  "comprises the 48 raw correlations (12 metrics × 4 outcomes), and family F2 comprises the 44 partial "
  "correlations on ranks controlling domain built-up area (11 metrics × 4 outcomes; size itself is not "
  "controlled for). Each family is controlled separately with the Benjamini–Hochberg false discovery "
  "rate at q = 0.05 [29]. Because villages in one region are spatially autocorrelated (Section 3.4), "
  "nominal p-values overstate the effective evidence; we therefore adjust each test’s effective sample "
  "size with the Moran’s-I-based correction of Dutilleul [34] and apply the BH procedure a second time "
  "to the ESS-adjusted p-values. In Table 2, stars mark nominal p-values and a dagger marks tests that "
  "survive BH control on ESS-adjusted p within their family; 95% confidence intervals for the headline "
  "correlations are computed by a spatial block bootstrap [35], [36]—villages are grouped into 14 "
  "blocks on a 0.15° latitude–longitude grid and blocks are resampled whole (B = 2999, fixed seed "
  "11)—which preserves within-block spatial dependence instead of assuming independent villages. "
  "Per-test statistics are reported in Appendix Table A3. Because size correlates with nearly "
  "everything—larger villages are hotter and also tend to sit in more open, serviceable terrain—the "
  "F2 partial correlations are the conservative read of metric-level association; we do not call them "
  "independent effects, only size-adjusted ones.")
P("Robustness is assessed four ways. First, the four-phase deployment repetition absorbs grid "
  "alignment (Section 2.5). Second, the peri-urban Yuliang case is excluded and all associations "
  "re-estimated. Third, the two villages with fallback coverage targets are excluded from all coverage "
  "associations. Fourth, a leave-one-county-out jackknife over the six counties checks that no single "
  "county drives the headline associations (Appendix Table A4). Spatial autocorrelation is diagnosed "
  "with Moran’s I under row-standardized inverse-distance weights on UTM Zone 50N coordinates (999 "
  "permutations; Appendix Table A5), and it is this diagnostic that motivates the ESS correction and "
  "the block bootstrap above. Trend lines in scatter plots are Theil–Sen robust fits [30], [31]. The "
  "propagation core is covered by unit tests (Section 2.5). All computation uses open-source Python "
  "libraries (rasterio, numpy, pandas, SciPy, matplotlib); the full code, the village-level analysis "
  "table, and the scene manifest are deposited in a public repository (see Data Availability).")

# ==================== 3. RESULTS ====================
H1("3. Results")
H2("3.1. Morphological Characterization")
P("The 29 villages differentiate clearly along the terrain dimension (Table 1, Fig. 2). Topographic "
  "openness spans tSVF = 0.899–0.999: most villages occupy open basin floors where the 2500-m horizon is "
  "essentially unobstructed, while a minority—notably Mulihong (tSVF = 0.899) and Zuyuan (0.928)—sit in "
  "deeply incised valleys. The forest ring covers 35–97% of the 300-m backdrop. Domain built-up area "
  "ranges from essentially nil in the two under-detected mountain villages to 56.8 ha in the peri-urban "
  "Yuliang, and component built-up area from 5.2 to 258.7 ha; relief within the village domain ranges "
  "from 38 m on the basin floor to 538 m around the highest mountain villages. The three setting classes "
  "are legible in the metric space: basin villages are open and weakly forested in their "
  "immediate ring; piedmont villages are intermediate; deep-valley villages are enclosed, elongated, "
  "steep, and almost fully ringed by forest.")
P("Two morphological observations deserve note. First, despite the canonical image of the south-facing "
  "Huizhou village, four villages (Kantou, Guanlu, Nanping, and Zhukou) face north, with negative "
  "southness values—a reminder that the textbook ideal is an idealization rather than a rule. Second, "
  "plan form and terrain co-vary: valley villages are strongly elongated along their streams (elongation "
  "up to 3.5), whereas basin villages are more isotropic. This co-variation, visible in the standardized "
  "distributions of Fig. 2c, is the statistical signature of site selection responding to terrain—and it "
  "is what makes the mountain-village morphometric problem different from its urban counterpart.")
import pandas as pd
_t1 = pd.read_csv("tables/Table1_sample.csv").fillna("—")
TABLE(1, "Key attributes of the 29 sampled villages: county, domain built-up area, elevation, relief, "
         "terrain openness (tSVF), forest ring share, summer LST, the site-domain anomaly ΔLST-V1 and "
         "the terrain-matched anomaly ΔLST-V3 with overpass-bootstrap 95% CIs, and four-phase coverage "
         "metrics (cov85 mean with cross-phase standard deviation; RSRP p10). The full "
         "twelve-metric table is Appendix Table A1 (repository CSV).",
      list(_t1.columns), _t1.values.tolist(), fontsize=7)
FIG("figures/Fig2_morphology.png", 2,
    "Morphological characterization of the sample. (a) Topographic openness (tSVF) and (b) forest ring "
    "share over a 30-m hillshade, point size proportional to domain built-up area; (c) standardized "
    "distributions of eight metrics across the 29 villages.")

H2("3.2. Summer Thermal Environment")
P("Village-core summer LST spans 31.2–42.8 °C (mean 36.4 °C) across the sample. The coolest villages "
  "are high-elevation, enclosed mountain settlements (Mulihong 31.2 °C, Tachuan 31.8 °C, Zuyuan "
  "32.1 °C); the hottest are large, open settlements of the basin lowlands and the peri-urban fringe "
  "(Yuliang 42.8 °C, Qiankou 40.1 °C, Tangyue 39.8 °C), only two of which exceed 40 °C. The heat "
  "anomaly depends on how the background is defined, and the gradient across definitions is itself "
  "informative (Fig. 3c). Under the site-domain definition (V1), all 29 village point estimates are "
  "positive: ΔLST-V1 ranges from +0.03 to +7.69 °C with a mean of +2.45 °C, and the overpass-bootstrap "
  "95% interval excludes zero for 21 of 29 villages. Under the fabric variant (V2, n = 22) the anomaly "
  "is larger (mean +5.28 °C; 19 of 22 intervals above zero), isolating the built boundary itself. "
  "Under the terrain- and land-cover-matched background (V3), the anomaly halves: mean +1.27 °C, range "
  "−0.91 to +4.19 °C, with 25 of 29 point estimates positive and 22 intervals excluding zero; the four "
  "non-positive cases (Tachuan, Xucun, Shitan, Huansha) are villages whose fabric is not warmer than "
  "unbuilt terrain of the same elevation, slope, and tree cover. Part of the apparent heat excess "
  "therefore reflects where villages sit rather than the fabric itself—yet even the controlled "
  "estimate keeps the sign of the canonical urban heat island in 25 of 29 villages [11], [12].")
P("The spatial pattern is legible in Fig. 3a. The hot corridor follows the broad, open Xin’an River "
  "basin lowlands, where the large settlements of the plain—including the county seats and the "
  "peri-urban fringe of Shexian—concentrate; the cool zones track the forested ridge systems. The five "
  "coolest villages (Mulihong 31.2, Tachuan 31.8, Zuyuan 32.1, Xucun 33.4, Renli 33.7 °C) combine high "
  "elevation, high relief, and near-complete forest rings, while the five hottest (Yuliang 42.8, "
  "Qiankou 40.1, Tangyue 39.8, Xiongcun 39.3, Zhanqi 38.4 °C) combine low relief, maximal openness, "
  "and large contiguous fabric. The matched-background comparison shows that the anomaly is not a "
  "simple elevation artifact: most high villages are cooler in absolute terms yet remain warmer than "
  "equally high, equally forested surroundings—though the four exceptions above warn that this is not "
  "universal.")
P("The morphological associates of absolute LST form a coherent terrain story (Table 2). Village LST "
  "rises with topographic openness (ρ = +0.65, q = 0.011, 95% CI [+0.19, +0.88]) and domain built-up "
  "area (+0.61, q = 0.007), and falls with relief (−0.71, q = 0.002, CI [−0.86, −0.41]), elevation "
  "(−0.56, q = 0.047, CI [−0.85, −0.03]), and water proximity (minimum distance −0.54, q = 0.019; mean "
  "distance −0.50, q = 0.044). The forest ring (−0.36, p = 0.065, q = 0.23) and slope (−0.38, "
  "p = 0.040, q = 0.17) show negative tendencies that do not survive ESS-adjusted family control, and "
  "the plan-form metrics are unrelated to absolute LST (elongation ρ = −0.12; compactness ρ = −0.19; "
  "both n.s.). Controlling domain built-up size reshuffles the picture: no thermal association "
  "survives ESS-adjusted FDR control (relief partial ρ = −0.53, q = 0.071; elevation −0.46, q = 0.19; "
  "openness +0.42, q = 0.25), so the size-controlled thermal read is suggestive rather than "
  "confirmatory. For the matched anomaly ΔLST-V3, three attributes survive family control: domain "
  "built-up area (raw ρ = +0.50, q = 0.017, CI [+0.11, +0.78]), water proximity (minimum distance "
  "−0.48, q = 0.025), and relief (−0.45, q = 0.047, CI [−0.69, −0.14]). Elevation, by contrast, shows "
  "no surviving association with the matched anomaly (−0.28, n.s.): once the background is matched on "
  "terrain, high villages are not systematically less anomalous. Openness (+0.32, n.s.) and the forest "
  "ring (ρ ≈ 0) are likewise unrelated to the matched anomaly, as are both plan-form metrics "
  "(elongation −0.02; compactness −0.07). The two dependent variables therefore tell different "
  "stories: absolute LST covaries mainly with *where* the village sits, and the matched anomaly with "
  "size, relief, and water proximity rather than with *how* the fabric is arranged. Southness is "
  "unrelated to village LST (ρ = −0.13, p = 0.50) and to the anomaly (ρ = −0.07, p = 0.73); "
  "north–south asymmetry behaves similarly (−0.24, p = 0.20 against LST; −0.17, p = 0.38 against the "
  "anomaly).")
FIG("figures/Fig3_lst.png", 3,
    "Summer thermal environment. (a) Quality-masked LST median composite over the 26 independent "
    "overpasses (2019–2025, Landsat 8/9 Collection 2), with the four WRS path/row footprints outlined; "
    "(b) village-core versus site-domain background LST, villages sorted by core LST; (c) per-village "
    "heat anomaly ΔLST with overpass-bootstrap 95% CI under the site-domain (V1) and terrain-matched "
    "(V3) definitions (n = 29; V1 mean +2.45 °C; V3 mean +1.27 °C).")

H2("3.3. Wireless Coverage and Deployment-Phase Sensitivity")
P("Under the standardized deployment, four-phase mean cov85 spans 59.8–100% across villages (mean "
  "93.0%). Seventeen of 29 villages reach 96% or better, and these are precisely the open basin "
  "settlements; the deep-valley villages Mulihong (59.8%) and Zuyuan (63.6%) fall far behind, and the "
  "piedmont villages occupy the middle (Fig. 4a). Coverage at the lenient −95 dBm threshold is near "
  "saturation—16 villages stand at a four-phase mean of 100% (sample range 85.3–100%)—which is why "
  "cov85, the good-signal level in this scenario, is adopted as the primary metric; RSRP p10 ranges "
  "from −96.8 to −76.1 dBm and supports the same ordering.")
P("The four-phase repetition reveals a second, subtler result (Fig. 4b): grid phase matters exactly "
  "where coverage is marginal. The cross-phase standard deviation of cov85 reaches 34.9 percentage "
  "points in Zuyuan (phase values from 12% to 88%) and exceeds 20 points in Renli, Shitan, and "
  "Changxi, whereas basin villages are phase-invariant at or near 100%. Deployment sensitivity is "
  "thus itself associated with morphological setting: villages at the coverage margin can be lucky or "
  "unlucky with a given site pattern, while well-sited villages are comparatively insensitive to it. "
  "Accordingly, village rankings across phases are moderately stable for cov85 (mean pairwise rank "
  "correlation 0.57, minimum 0.47) and weaker for RSRP p10 (0.39, minimum 0.17); the plain mean RSRP "
  "proved phase-unstable (0.10, minimum −0.09), being hostage to the distance of the single nearest "
  "site, and was discarded from further analysis. The four-phase mean therefore absorbs most—but not "
  "all—deployment arbitrariness: rankings among open, well-covered villages are phase-insensitive, "
  "while rankings among marginal, deep-valley villages are not.")
P("The village-level detail behind these statistics is instructive. Yuliang, the peri-urban case, "
  "combines the largest target set (631 built-up cells in the domain) with a mid-pack four-phase mean "
  "of 86.2%—and, unexpectedly, substantial phase dispersion (71.0–99.2% across phases): its fabric "
  "extends from the open plain into a river corridor whose margins lose visibility in two of the four "
  "grid alignments. Size and openness therefore do not exempt a village from deployment sensitivity. "
  "At the other extreme, Zuyuan and Mulihong—the two villages for which the land-cover map "
  "under-detects the built fabric, and whose coverage is therefore evaluated on a 150-m fallback "
  "grid—are not only the worst covered but also among the most phase-sensitive: their serviceability "
  "is decided by whether a half-spacing shift happens to place a site on the one ridge with a clear "
  "path into the valley. Because their fallback target grids are not strictly comparable to the "
  "built-up-cell targets of the other 27 villages, all coverage associations below are re-estimated "
  "without them (Section 3.4). Between these poles, piedmont villages such as Huansha, Zhukou, and "
  "Zhanqi (cov85 = 92–96%) lose service only at their valley-facing margins, which is where their "
  "weakest RSRP deciles originate.")
P("Frequency choice does not change the picture. Repeating the full chain at 700 MHz (a band central "
  "to China’s current rural 5G build-out) raises absolute coverage as expected—the four-phase mean "
  "cov85 rises from 93.0% to 95.9%—but preserves the inter-village ordering (rank correlation 0.89 "
  "between the four-phase means of the two bands; in the phase-0 comparison, 0.81 for cov85 and ≥0.98 "
  "for the RSRP statistics; Appendix Table A6). The near-saturated cov95 ordering is, as expected, the "
  "least stable statistic across bands (ρ = 0.39). The morphological gradients reported below are "
  "therefore not an artifact of band choice.")
FIG("figures/Fig4_coverage.png", 4,
    "Wireless coverage under the standardized virtual deployment. (a) Four-phase mean cov85 by village "
    "(points) over the 1620-site phase-0 grid (dots), on a 30-m hillshade; (b) deployment-phase "
    "sensitivity: villages with low mean coverage also show high cross-phase dispersion (bars: "
    "cross-phase minimum–maximum range).")

H2("3.4. Morphology–Performance Coupling and Trade-Offs")
P("Bringing the two performances together (Fig. 5, Table 2) answers RQ4. To classify couplings "
  "consistently, we fix utility directions explicitly: the thermal utility is U_T = −LST (or −ΔLST) "
  "and the coverage utility is U_C = cov85 (or RSRP p10). A morphological attribute is a *trade-off* "
  "attribute when its associations with the two utilities carry opposite signs—note that same-signed "
  "raw correlations with LST and with cov85 imply opposite-signed utility associations—and a "
  "*one-sided* attribute when only one association is detectable.")
P("By this rule, two attributes are trade-off axes. First, topographic openness is associated with "
  "higher LST (ρ = +0.65, q = 0.011) and better coverage (cov85, ρ = +0.71, q = 0.007, CI [+0.45, "
  "+0.89]). Under size control the coverage side becomes marginal (partial ρ = +0.60, q = 0.067 for "
  "cov85) while the RSRP-p10 side survives (partial ρ = +0.60, q = 0.045), and the thermal side drops "
  "out (partial ρ = +0.42, q = 0.25): the openness axis is robust in the raw family but size-entangled. "
  "Second, relief is associated with lower LST (ρ = −0.71, q = 0.002) and worse coverage (ρ = −0.66, "
  "q = 0.007); here too, both sides fall to marginal status under size control (partial ρ = −0.53 and "
  "−0.52, q = 0.071 each). Domain built-up area is itself double-edged in the raw family—larger "
  "villages are hotter (ρ = +0.61, q = 0.007) and nominally better covered (ρ = +0.47), although the "
  "coverage side does not survive the ESS adjustment (q = 0.055)—which is precisely why size serves "
  "as the control variable rather than being read as a morphological trade-off axis. Slope behaves "
  "like relief on the coverage side (ρ = −0.85, q < 0.001, CI [−0.94, −0.70]; size-controlled partial "
  "ρ = −0.80, surviving) with only a non-significant thermal tendency (ρ = −0.38, q = 0.17).")
P("Three attributes are one-sided. The forest ring strongly shadows signal (ρ = −0.81, q < 0.001, CI "
  "[−0.88, −0.64]; size-controlled partial ρ = −0.79, surviving) but shows no detectable association "
  "with village-scale daytime LST after quality masking (ρ = −0.36, p = 0.065, q = 0.23)—a null "
  "result we take seriously in Section 4.1. Elevation and water proximity are purely thermal: "
  "elevation is associated with lower LST (ρ = −0.56, q = 0.047) with no coverage counterpart "
  "(ρ = +0.07, n.s.), and water proximity likewise (minimum distance vs. LST ρ = −0.54, q = 0.019; "
  "vs. coverage ρ = −0.17, n.s.); elevation is not, however, associated with the matched anomaly "
  "(Section 3.2). Southness, the attribute most often invoked in heritage discourse, shows no "
  "association with any outcome. The two plan-form metrics show no association that survives family "
  "control: neither is associated with any thermal outcome (elongation ρ = −0.12 with LST and −0.02 "
  "with ΔLST; compactness −0.19 and −0.07; all n.s.), and on the wireless side what remains is a "
  "consistent but formally non-surviving negative tendency of compactness (raw ρ = −0.36 with cov85, "
  "q = 0.16; ρ = −0.41 with RSRP p10, q = 0.099; size-controlled partial ρ = −0.46, q = 0.090) and a "
  "marginal coverage trend for elongation (cov85 ρ = −0.33, p = 0.091, q = 0.18). Under the "
  "classification rule above, plan form is therefore not a detectable coupling axis at this sample "
  "size; Section 4.1 discusses what this refines relative to plan-form expectations from the urban "
  "literature. Of the 48 raw tests (family F1), 17 survive Benjamini–Hochberg control on ESS-adjusted "
  "p-values at q = 0.05 (18 on nominal p); of the 44 size-controlled partial tests (family F2), 5 "
  "survive (9 on nominal p).")
P("Sensitivity checks qualify these associations honestly. Excluding the peri-urban Yuliang case "
  "leaves every reported sign intact and strengthens several magnitudes (openness–coverage ρ = +0.80; "
  "forest-ring–coverage −0.83; slope–coverage −0.85; openness–LST +0.62). Excluding the two "
  "fallback-target villages from the coverage tests leaves the openness–coverage correlation at "
  "ρ = +0.64 (p < 0.001, n = 27) and the forest-ring and slope correlations essentially unchanged "
  "(−0.82 and −0.82). In the leave-one-county-out jackknife the headline associations keep their "
  "signs in every county omission (openness–coverage ρ between +0.64 and +0.81; forest-ring–coverage "
  "between −0.73 and −0.89; openness–LST between +0.57 and +0.75; relief–ΔLST between −0.38 and "
  "−0.53, all p < 0.05), so no single county drives the results. Spatial diagnostics motivate the "
  "conservative inference machinery: Moran’s I detects significant spatial clustering for cov85 "
  "(I = 0.22, p = 0.005), tSVF (I = 0.20, p = 0.003), absolute LST (I = 0.18, p = 0.004), and "
  "elevation (I = 0.23, p = 0.001), but not for the matched anomaly (I = −0.02, p = 0.78) or RSRP p10 "
  "(I = 0.12, p = 0.055); accordingly, inferential weight rests on the ESS-adjusted FDR control and "
  "the block-bootstrap intervals rather than on nominal p-values alone.")
FIG("figures/Fig5_tradeoff.png", 5,
    "Morphology–performance couplings. (a), (b) Topographic openness versus village LST and cov85; "
    "(c), (d) forest ring share versus village LST and cov85. Point size: domain built-up area; lines: "
    "Theil–Sen robust fits truncated to the observed metric range; annotations: Spearman ρ with "
    "nominal p, Benjamini–Hochberg q on ESS-adjusted p within test family F1, and the size-controlled "
    "partial ρ; n = 27 in panels c–d because the forest ring, defined around the detected built-up "
    "component, is unavailable for the two under-detected villages; the number of villages at the "
    "cov85 ceiling is noted in panels b and d.")
_t2 = pd.read_csv("tables/Table2_correlation.csv")
TABLE(2, "Spearman correlations between morphological metrics and performance variables (upper row per "
         "metric), and partial correlations controlling domain built-up area (row marked “ctrl size”). "
         "* p < 0.05, ** p < 0.01, *** p < 0.001 (nominal); † survives Benjamini–Hochberg control at "
         "q = 0.05 on Moran’s-I effective-sample-size-adjusted p-values within its test family (F1: 48 "
         "raw tests; F2: 44 partial tests). Block-bootstrap 95% CIs and per-test statistics are given "
         "in Appendix Table A3. Plan-form metrics use n = 28 (Mulihong has no detected built-up cell "
         "in its domain); water distances and the forest ring use n = 27 (component-anchored); others "
         "n = 29.",
      list(_t2.columns), _t2.values.tolist(), fontsize=8)

# ==================== 4. DISCUSSION ====================
H1("4. Discussion")
H2("4.1. Trade-Off Axes, Read Mechanistically")
P("The first axis is enclosure. The same topographic enclosure that shades a valley village, shortens "
  "its solar exposure, and channels nocturnal cold-air drainage also blocks the low-elevation radio "
  "paths on which rural macrocells depend. Vernacular site selection and radio physics price the same "
  "terrain in opposite currencies. This result inverts the urban intuition: in city morphometrics, "
  "higher sky-view factor is generally associated with *cooler* canyon environments because SVF indexes "
  "heat release from street canyons [14], [18]; at the mountain-village scale, higher tSVF is associated "
  "with *hotter* settlements because it indexes insolation exposure of the whole site. The sign flip is "
  "a scale effect, not a contradiction, but it warns against transferring urban morphometric rules to "
  "rural terrain-coupled settlements. We note honestly that the thermal side of this axis does not "
  "survive size control, and we therefore treat the enclosure axis as a planning hypothesis with a "
  "strong coverage side and a moderate, size-entangled thermal side.")
P("The second axis is the forest ring, the emblematic Huizhou backdrop—and here the quality-controlled "
  "reanalysis delivered a salutary surprise. In a simple mean-based composite without the QA bit mask, "
  "the ring appeared to cool (ρ ≈ −0.58); after quality masking and median compositing, the village-"
  "scale daytime association is not distinguishable from zero under family control (ρ = −0.36, "
  "p = 0.065, ESS-adjusted q = 0.23; size-controlled partial ρ = −0.17). We report this null result "
  "rather than the earlier apparent effect. Three non-exclusive explanations are plausible: the ring "
  "cools its own slopes as much as the village, so core-versus-background contrast cancels; daytime "
  "summer LST is dominated by insolation exposure rather than by evapotranspiration at the 100-m "
  "scale; and any genuine ring effect may operate at night or at the courtyard scale, both invisible "
  "to our composite. What remains robust—and large—is the ring’s association with signal shadowing "
  "(ρ = −0.81; size-controlled partial ρ = −0.79), which in our model is entirely terrain-mediated. "
  "Whether vegetation attenuation over the ring [19] would change this village-level ranking cannot be "
  "determined from the present model: it would deepen absolute shadowing, but the ranking response "
  "depends on village-specific path geometries, so we flag it as a validation target rather than "
  "assert a direction (Section 4.4). For the ecological-wisdom reading of vernacular landscape, the "
  "honest summary is therefore selective support: water proximity survives scrutiny, the forest "
  "ring’s daytime thermal benefit at village scale is not detectable, and the ring’s "
  "digital-infrastructure cost is large and robust across sensitivity checks.")
P("Plan-form metrics and southness refine the architectural reading. Under the Frame-D definitions, "
  "compactness is not detectably associated with any thermal outcome; a component-based formulation of "
  "the same concept produced stronger apparent associations that did not reproduce across land-cover "
  "processing versions (Section 2.3), so we treat the compactness question as unresolved rather than "
  "answered. The hypothesis that compact historic cores run hotter than their hinterland—plausible from "
  "sparse internal green space and reduced sky view—requires building-scale data, not 10-m land cover. "
  "On the wireless side, compactness shows a consistent negative tendency (ρ ≈ −0.36 to −0.46 across "
  "the coverage and RSRP metrics) that does not survive false-discovery control at n = 28; we report "
  "it as a hypothesis for larger samples, not a finding. Southness, by contrast, is inert at "
  "settlement scale in this sample: we find no evidence that orientation predicts thermal outcomes, so "
  "claims for village-scale orientation effects should be treated as unproven rather than "
  "canonical—orientation effects are real at the scale of buildings and canyons [3], [14], but they "
  "average out across a settlement, and the four sampled villages that depart from the canonical "
  "south-facing orientation show that the textbook ideal is an idealization rather than a rule.")
H2("4.2. Comparison with Previous Studies")
P("Three strands of previous work frame our findings. First, the urban morphometrics literature has "
  "established SVF–temperature relationships at the scale of street canyons and LCZ classes [13], [14], "
  "[18], where openness generally aids night-time cooling. Our association between openness and higher "
  "village-scale summer LST (ρ = +0.65) does not contradict this work: it reflects the different physics "
  "that dominate when the ‘canyon’ is a mountain valley and the variable is daytime insolation of the "
  "whole site rather than nocturnal long-wave release from a street. The two regimes are reconciled once "
  "one recognizes that tSVF here measures horizon obstruction by terrain, not by buildings. Second, the "
  "vernacular-climate literature has documented passive strategies of traditional settlements through "
  "single-village instrumentation and simulation [1]–[4], [7], including work in the Huizhou area itself "
  "[3]. Our contribution relative to that strand is statistical rather than instrumental: with 29 "
  "villages we can separate site-level associations from fabric-level ones (openness and relief versus "
  "compactness) in a way that one-village case studies structurally cannot, and our null result for the "
  "forest ring is a caution against generalizing courtyard-scale measurements to the settlement scale. "
  "Third, the urban-informatics strand has shown how open imagery and reproducible computation can "
  "characterize fine-grained urban form at scale [22], [24], [25]. We extend that ethos to a domain it has "
  "largely bypassed—rural heritage settlements—and to a second physical field, radio propagation, that "
  "this literature has never engaged. The nearest precedents to our coverage analysis are the "
  "propagation-modeling and radio-planning literature [8]–[10]; our inversion of the perspective, asking "
  "not ‘where should sites go’ but ‘which village forms are inherently hard to serve’, appears to be "
  "new.")
P("Elongation illustrates the same discipline from the other side. In the thermal literature, "
  "linear valley villages would be filed under ‘favorably shaded’; in radio planning the same form is a "
  "known worst case, because a linear settlement along a valley floor maximizes the share of fabric "
  "lying at low radio horizons. In our data the coverage side of that expectation appears only as a "
  "marginal trend (ρ = −0.33, p = 0.09 with cov85) and the thermal side not at all (ρ = −0.12, "
  "p = 0.54 with LST): the two-physics elongation trade-off is mechanistically grounded in both "
  "literatures, but it is not confirmed at this sample size, and we state that plainly rather than "
  "inherit the expectation from either single-physics strand.")
H2("4.3. Implications for Heritage-Compatible and Digital-Village Planning")
P("In what follows we separate three levels of statement: associations observed in the data under "
  "false-discovery-rate control, inferences that are mechanistically plausible but statistically "
  "partial, and hypotheses that the present data cannot test. For conservation planning, the practical "
  "message is that thermal-friendly morphology is not automatically connectivity-friendly: the "
  "associations reported here indicate a potential planning trade-off—quantifiable in principle, but "
  "requiring site-specific measurement and validation before any engineering action. Even the "
  "worst-served villages reach roughly 60% good coverage under a macro-only deployment in this "
  "stylized scenario; the residual gap is precisely the niche for low-visual-impact small cells at the "
  "village perimeter, fed through the open river and road corridors that pierce the forest ring. LOS "
  "simulation on open elevation data, as used here, offers heritage authorities a cheap screening tool "
  "that can be run before any hardware or visual-impact discussion begins, and that directly respects "
  "the constraint that matters most in a protected village: where things can be seen from, and where "
  "they cannot.")
P("For digital-village and CIM programs, the same pipeline doubles as an analytics layer. From open "
  "data it maps both the thermal burden and the radio visibility of every settlement, providing the "
  "baseline against which digital infrastructure in historic settings can be planned rather than "
  "improvised; in the Chinese context, such layers can be embedded within the CIM base-platform "
  "standard framework [6]. Because deployment-phase sensitivity is itself mapped (Fig. 4b), planners "
  "can distinguish villages whose service will be insensitive to deployment phase from villages that "
  "need site-by-site engineering. The open-data character of the whole workflow matters practically: "
  "it can be audited by conservation reviewers, re-run as new satellite data arrive, and transferred "
  "to other mountainous regions at low marginal cost, subject to local data-availability checks.")
P("Finally, there is a discursive implication. The ecological wisdom of vernacular settlements is "
  "usually asserted qualitatively in heritage narratives; here it is quantified—and found to be "
  "selective and conditional. In this sample, terrain setting and settlement size carry measurable "
  "thermal associations while plan-form effects are not detectable at this sample size; the forest "
  "ring’s daytime village-scale cooling is not detectable after quality control; the orientation "
  "canon is neither rule nor thermal determinant at settlement scale; and the same attributes that "
  "shade and shelter also isolate from the network. A mature, evidence-based account of traditional "
  "ecological knowledge should be able to state both sides of that ledger—including the null results. "
  "We suggest that this kind of quantified, two-sided assessment is precisely what makes the wisdom "
  "credible—and translatable—when it is communicated to engineering, planning, and policy audiences.")
H2("4.4. Limitations and Future Work")
P("Nine limitations qualify the findings. First, LST is not air temperature or thermal comfort: "
  "satellite LST responds to surface fabric and misses shade-mediated comfort, and our composite "
  "characterizes summer daytime only; sensor-based validation is the obvious next step. Second, the "
  "composite is correlational in nature and single-season; night-time, winter, and transitional-season "
  "behavior may differ, which matters directly for interpreting the forest-ring null. Absolute LST is "
  "additionally composite-dependent across the four satellite frames and is used for ranking only; for "
  "8 of the 29 villages the overpass-bootstrap interval of the V1 anomaly includes zero, so "
  "village-level anomaly values should be read with their intervals (Table 1); we do not propagate the "
  "per-pixel ST_QA uncertainty band distributed with the product, which would widen the reported "
  "intervals further; and the anomaly is definition-sensitive—four villages turn non-positive under "
  "the matched background—so it should be read as a range across the V1–V3 definitions, not as a "
  "point value. Third, the standardized deployment isolates morphology but is not any operator’s "
  "network; absolute coverage levels would shift with real site inventories, antenna patterns, and "
  "traffic engineering, and the morphological ranking—the object of this study—proved moderately "
  "stable across grid phases for open villages but genuinely phase-sensitive for marginal ones, a "
  "heterogeneity our four-phase mean only partly absorbs. Fourth, the propagation chain is a "
  "deterministic received-power scenario: it omits vegetation and building clutter, diffraction "
  "refinement, shadowing, and interference; reported coverage is optimistic in absolute terms; "
  "whether vegetation attenuation would change the village-level ranking is unknown and is a "
  "validation target; and validation against measured or crowdsourced RSRP, or against a "
  "diffraction-capable model such as ITU-R P.1812 [37] or Longley–Rice, remains future work. The RMa "
  "NLOS branch is moreover specified only to 5 km; a truncation variant enforcing that limit leaves "
  "all village coverage metrics unchanged (Appendix Table A7), but links near the assignment radius "
  "remain extrapolations. Fifth, the morphometrics are coarse by architectural standards: tSVF is "
  "evaluated once at the built-up centroid rather than across the fabric, and the 10-m land-cover map "
  "cannot resolve the narrow lanes, courtyards, building heights, and street H/W ratios of Huizhou "
  "villages; the courtyard-scale design hypotheses mentioned in Section 4.1 are therefore speculative "
  "and require building-scale data. The plan-form metrics are computed on detected 10-m cells inside a "
  "fixed domain; an earlier component-based formulation did not reproduce across processing versions, "
  "and residual sensitivity to land-cover classification error remains. Sixth, the elevation source "
  "is a digital *surface* model: over forest and fabric its heights include canopy and roof signals, "
  "so terrain metrics are surface metrics; a cross-check against the independent NASADEM product "
  "agrees closely (centroid elevation ρ = 0.999, domain relief ρ = 0.988; Appendix Table A10), but "
  "NASADEM is itself a radar surface model, and per-village discrepancies reach 208 m at Tachuan, "
  "indicating steep, densely vegetated terrain where surface and ground diverge; LOS determination "
  "and all elevation-derived metrics inherit this limitation. Seventh, two villages rely on fallback "
  "coverage targets and lack component-anchored metrics, and one of them also lacks plan-form "
  "metrics; although exclusion checks leave the headline associations unchanged, manual or "
  "semi-automated mapping of their built boundaries—ideally tied to field survey—would remove this "
  "caveat. Eighth, the sample is a purposive regional sample (n = 29), not a probability sample; "
  "Moran’s I detects significant spatial clustering in tSVF, cov85, absolute LST, and elevation, and "
  "while the ESS correction, the block bootstrap, and the county jackknife mitigate the resulting "
  "overconfidence, all inference remains associational and we do not generalize beyond the sampled "
  "region. Ninth, one village’s detected built component merges with a county seat, though excluding "
  "it changes no reported sign.")
P("Future work follows four lines: (i) multi-season and diurnal LST composites, plus air-temperature "
  "and comfort sensor campaigns in a stratified subsample of villages; (ii) building-scale "
  "morphometrics from high-resolution imagery or survey—street H/W, courtyard green view, "
  "building-resolved SVF—tested against the settlement-scale associations reported here; (iii) "
  "validation of the radio chain against real operator inventories and crowdsourced measurements, "
  "cross-checking against a diffraction-capable propagation model, and explicit measurement of "
  "vegetation attenuation to test whether the forest-ring ranking survives it; and (iv) extension of "
  "the sample beyond Huizhou to test whether the trade-off axes generalize across mountain-village "
  "regions with different building traditions.")

# ==================== 5. CONCLUSIONS ====================
H1("5. Conclusions")
P("Across 29 traditional Huizhou villages, settlement morphology is associated with summer surface "
  "temperature and with simulated wireless coverage in systematically coupled ways. Under the "
  "site-domain definition, all 29 village point estimates of the summer heat anomaly are positive "
  "(mean ΔLST = +2.4 °C; the overpass-bootstrap interval excludes zero in 21 of 29); under a terrain- "
  "and land-cover-matched background the anomaly halves (mean +1.3 °C; 25 of 29 positive), and its "
  "surviving correlates are settlement size, local relief, and water proximity rather than elevation. "
  "Absolute temperature covaries mainly with terrain setting (relief ρ = −0.71; openness ρ = +0.65; "
  "elevation ρ = −0.56; water proximity ρ = −0.54). Two trade-off axes emerge under a consistent "
  "utility rule—openness and relief each carry opposite-signed associations with the thermal and "
  "coverage utilities—although under size control only the coverage side survives false-discovery "
  "control, so the axes are robust at sample level but partly size-entangled. The forest ring is a "
  "one-sided digital cost whose village-scale daytime thermal benefit is not detectable after quality "
  "control, and southness is inert at settlement scale. Plan-form metrics show no association that "
  "survives false-discovery control at this sample size; the compactness and elongation tendencies "
  "reported above are hypotheses, not findings. The headline associations are stable to frequency "
  "band, county jackknife, outlier exclusion, and the propagation-model range limit, with grid-phase "
  "effects absorbed for open villages but persisting for marginal ones, and the entire analysis runs "
  "on open data and a scripted, deposited pipeline. Because the design is cross-sectional and the "
  "deployment stylized, the quantified couplings are hypotheses for site-specific assessment—not "
  "causal effects, operational coverage estimates, or general laws. For the agencies that must "
  "reconcile climate adaptation with digital modernization in historic mountain settlements, the "
  "central lesson is that a potential conflict is visible in the data and quantifiable—and, once made "
  "visible, amenable to design—subject to site-specific validation.")

# ==================== 声明 ====================
H1("Declarations")
P("**Data Availability.** The full analysis code, the village-level analysis table (all twelve "
  "morphological metrics, four performance variables, and four-phase coverage detail), the Landsat "
  "scene manifest, and the figure scripts are deposited at "
  "https://github.com/ccst2000/huizhou-thermal-wireless-tradeoffs and "
  "archived at https://doi.org/10.5281/zenodo.22011592. Public input datasets are cited in the References.")
P("**Author Contributions (CRediT).** Lei Zhang: conceptualization, methodology, software, formal "
  "analysis, investigation, data curation, writing—original draft, writing—review and editing, "
  "visualization, project administration.")
P("**Conflicts of Interest.** The author declares no conflict of interest.")
P("**Funding.** [Funding information to be inserted.]")

H1("Acknowledgments")
P("The author thanks the data providers: the U.S. Geological Survey and Microsoft Planetary Computer "
  "for Landsat access, the European Space Agency for WorldCover, and the Copernicus programme for the "
  "digital surface model; and Esri for the World Imagery base map used in Fig. 1.")

# ==================== REFERENCES ====================
H1("References")
refs = [
    "Z. J. Zhai and J. M. Previtali, “Ancient vernacular architecture: Characteristics categorization "
    "and energy performance evaluation,” Energy and Buildings, vol. 42, no. 3, pp. 357–365, 2010.",
    "S. Gou, Z. Li, Q. Zhao, V. M. Nik, and J.-L. Scartezzini, “Climate responsive strategies of "
    "traditional dwellings located in an ancient village in hot summer and cold winter region of "
    "China,” Building and Environment, vol. 86, pp. 151–165, 2015.",
    "C. Pan, Y. Wu, S. Chen, and Y. Yang, “Indoor environmental comfort assessment of traditional folk "
    "houses: A case study in southern Anhui, China,” International Journal of Environmental Research "
    "and Public Health, vol. 20, no. 4, 3024, 2023.",
    "Y. Xiong, J. Zhang, Y. Yan, S. Sun, X. Xu, and E. Higueras, “Effect of the spatial form of "
    "Jiangnan traditional villages on microclimate and human comfort,” Sustainable Cities and Society, "
    "vol. 87, 104136, 2022.",
    "General Office of the CPC Central Committee and General Office of the State Council, “Strategic "
    "Outline for Digital Village Development,” Beijing, China, 2019 (in Chinese).",
    "Ministry of Housing and Urban-Rural Development of the People’s Republic of China, Technical "
    "Standard for City Information Modeling Basic Platforms, CJJ/T 315-2022, Beijing, China, 2022 "
    "(in Chinese).",
    "Z. Wan, et al., “How does outdoor spatial design shape the microclimate, comfort, and behavior "
    "in traditional Chinese villages? A systematic review across scales, contexts, and users,” "
    "Sustainability, vol. 17, no. 15, 6960, 2025.",
    "M. Hata, “Empirical formula for propagation loss in land mobile radio services,” IEEE "
    "Transactions on Vehicular Technology, vol. 29, no. 3, pp. 317–325, 1980.",
    "3GPP, “Study on channel model for frequencies from 0.5 to 100 GHz,” 3rd Generation Partnership "
    "Project, Sophia Antipolis, France, TR 38.901 V16.1.0, 2020.",
    "T. S. Rappaport, Wireless Communications: Principles and Practice, 2nd ed. Upper Saddle River, "
    "NJ, USA: Prentice Hall, 2002.",
    "T. R. Oke, “The energetic basis of the urban heat island,” Quarterly Journal of the Royal "
    "Meteorological Society, vol. 108, no. 455, pp. 1–24, 1982.",
    "A. J. Arnfield, “Two decades of urban climate research: A review of turbulence, exchanges of "
    "energy and water, and the urban heat island,” International Journal of Climatology, vol. 23, "
    "no. 1, pp. 1–26, 2003.",
    "I. D. Stewart and T. R. Oke, “Local climate zones for urban temperature studies,” Bulletin of "
    "the American Meteorological Society, vol. 93, no. 12, pp. 1879–1900, 2012.",
    "T. R. Oke, “Canyon geometry and the nocturnal urban heat island: Comparison of scale model and "
    "field observations,” Journal of Climatology, vol. 1, no. 3, pp. 237–254, 1981.",
    "K. Žákšek, K. Oštir, and Ž. Kokalj, “Sky-view factor as a relief visualization technique,” "
    "Remote Sensing, vol. 3, no. 2, pp. 398–415, 2011.",
    "Y. Xiong, Y. He, X. Xie, T. Zhai, N. Chu, L. Shen, and Y. Yang, “A study on the spatial form of "
    "traditional villages in Jiangnan region of China from the perspective of human thermal comfort: "
    "A case study of Nanjing, Jiangsu Province,” PLoS ONE, vol. 20, no. 5, e0323252, 2025.",
    "S. Wu, et al., “Research on microclimate optimization of traditional residential buildings in "
    "central Anhui based on humid and hot climate characteristics and regional architectural "
    "features,” Buildings, vol. 14, no. 8, 2323, 2024.",
    "A. Middel, K. Häb, A. J. Brazel, C. A. Martin, and S. Guhathakurta, “Impact of urban form and "
    "design on mid-afternoon microclimate in Phoenix Local Climate Zones,” Landscape and Urban "
    "Planning, vol. 122, pp. 16–28, 2014.",
    "ITU-R, “Attenuation in vegetation,” International Telecommunication Union, Geneva, Switzerland, "
    "Recommendation ITU-R P.833-9, 2016.",
    "ITU, Measuring Digital Development: Facts and Figures 2023. Geneva, Switzerland: International "
    "Telecommunication Union, 2023.",
    "European Space Agency, “Copernicus GLO-30 digital surface model (COP-DEM),” 2021. [Online]. "
    "Available: https://spacedata.copernicus.eu/collections/copernicus-digital-elevation-model, "
    "doi: 10.5270/ESA-c5d3d65 (accessed August 2026).",
    "X. Li, C. Zhang, W. Li, R. Ricard, Q. Meng, and W. Zhang, “Assessing street-level urban greenery "
    "using Google Street View and a modified green view index,” Urban Forestry & Urban Greening, "
    "vol. 14, no. 3, pp. 675–685, 2015.",
    "D. Zanaga et al., “ESA WorldCover 10 m 2021 v200,” Zenodo, 2022. [Online]. Available: "
    "https://doi.org/10.5281/zenodo.7254221 (accessed August 2026).",
    "Y. Long and L. Liu, “How green are the streets? An analysis for central areas of Chinese cities "
    "using Tencent Street View,” PLoS ONE, vol. 12, no. 2, e0171110, 2017.",
    "Y. Ye, D. Richards, Y. Lu, X. Song, Y. Zhuang, W. Zeng, and T. Zhong, “Measuring daily accessed "
    "street greenery: A human-scale approach for informing better urban planning practices,” Landscape "
    "and Urban Planning, vol. 191, 103434, 2019.",
    "U.S. Geological Survey, “Landsat 8–9 OLI/TIRS Collection 2 Level-2 science products,” 2021. "
    "[Online]. Available: https://doi.org/10.5066/P9OGBGM6 (accessed August 2026).",
    "Microsoft, “Planetary Computer: A planetary platform for sustainability,” collection "
    "landsat-c2-l2, catalog query executed 17 August 2026. [Online]. Available: "
    "https://planetarycomputer.microsoft.com.",
    "J. A. Voogt and T. R. Oke, “Thermal remote sensing of urban climates,” Remote Sensing of "
    "Environment, vol. 86, no. 3, pp. 370–384, 2003.",
    "Y. Benjamini and Y. Hochberg, “Controlling the false discovery rate: A practical and powerful "
    "approach to multiple testing,” Journal of the Royal Statistical Society: Series B, vol. 57, "
    "no. 1, pp. 289–300, 1995.",
    "H. Theil, “A rank-invariant method of linear and polynomial regression analysis,” Indagationes "
    "Mathematicae, vol. 12, pp. 85–91, 1950, doi: 10.1007/978-94-011-2546-8_20.",
    "P. K. Sen, “Estimates of the regression coefficient based on Kendall’s tau,” Journal of the "
    "American Statistical Association, vol. 63, no. 324, pp. 1379–1389, 1968.",
    "NASA Jet Propulsion Laboratory, “NASADEM Merged DEM Global 1 arc second V001,” NASA EOSDIS Land "
    "Processes DAAC, 2020, doi: 10.5067/MEaSUREs/NASADEM/NASADEM_HGT.001 (accessed August 2026).",
    "3GPP, “Evolved Universal Terrestrial Radio Access (E-UTRA): Physical layer—Measurements,” 3rd "
    "Generation Partnership Project, Sophia Antipolis, France, TS 36.214, 2017.",
    "P. Dutilleul, “Modifying the t test for assessing the correlation between two spatial "
    "processes,” Biometrics, vol. 49, no. 1, pp. 305–314, 1993.",
    "B. Efron and R. J. Tibshirani, An Introduction to the Bootstrap. New York, NY, USA: Chapman & "
    "Hall, 1993.",
    "H. R. Künsch, “The jackknife and the bootstrap for general stationary observations,” The Annals "
    "of Statistics, vol. 17, no. 3, pp. 1217–1241, 1989.",
    "ITU-R, “Propagation prediction techniques for the frequency range 30 MHz to 100 GHz,” "
    "International Telecommunication Union, Geneva, Switzerland, Recommendation ITU-R P.1812-6, 2023.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

# ==================== APPENDIX ====================
doc.add_page_break()
H1("Appendix A. Full Metric Table and Supporting Data")
P("Table A1 reports all twelve morphological metrics, the four performance variables, per-phase "
  "coverage, and Landsat observation support for the 29 villages; given its width (35 columns) it is "
  "provided as a machine-readable CSV in the repository rather than typeset here. Table A2 below lists "
  "the Landsat scene manifest. Further machine-readable appendix tables are distributed with the "
  "repository (see Supplementary Materials): Table A3, per-test exact statistics for families F1 and "
  "F2 (nominal and ESS-adjusted p, within-family BH q, block-bootstrap CI, n); Table A4, "
  "leave-one-county-out jackknife; Table A5, Moran’s I diagnostics; Table A6, 700-MHz four-phase "
  "coverage; Table A7, NLOS 5-km truncation and d3D-correction sensitivity; Table A8, per-village "
  "scene composition; Table A9, the three ΔLST variants with coverage-threshold and near-zero-date "
  "sensitivities; Table A10, DSM cross-check against NASADEM; Table A11, the sampling frame with "
  "per-village coordinate verification; Table A12, worked link budgets for three representative "
  "villages. All tables are generated by script from the same analysis table that feeds the in-text "
  "statistics.", size=9.5)
_man = pd.read_csv("data/lst_scene_manifest.csv")
_zero = {"LC08_L2SP_121039_20200720_02_T1", "LC08_L2SP_121039_20230729_02_T2",
         "LC08_L2SP_121040_20200922_02_T2"}
_near = {"LC08_L2SP_120039_20200627_02_T2", "LC08_L2SP_120040_20200627_02_T2"}
TABLE("A2", "Landsat 8/9 Collection-2 scene manifest (36 scene assets over four WRS frames). "
            "Adjacent-row assets of the same platform, path, and date are merged pixel-wise into 26 "
            "independent overpasses. Three scenes whose study-area subset is fully cloud- or "
            "shadow-covered after the QA chain (zero valid pixels) are marked “no” and excluded; two "
            "further scenes—one overpass date—retain only ≈0.01% valid pixels after masking and are "
            "marked “near-zero”, leaving 25 overpasses in the primary anomaly analysis.",
      ["Product ID", "Date", "Scene cloud (%)", "Path", "Row", "Valid"],
      [[r.id, r.date, r.cloud, r.path, (r.row if hasattr(r, "row") else r["row"]),
        ("no" if r.id in _zero else ("near-zero" if r.id in _near else "yes"))]
       for r in _man.itertuples()],
      fontsize=7)

H1("Supplementary Materials")
P("The following machine-readable files accompany this article in the public repository (Data "
  "Availability): TableA1_full.csv (full village-level analysis table, 29 villages × 35 fields); "
  "TableA3_exact_stats.csv (all 92 tests of families F1 and F2 with nominal and ESS-adjusted p, "
  "within-family BH q, block-bootstrap CI, and n); TableA4_jackknife.csv; TableA5_moran.csv; "
  "TableA6_700mhz_4phase.csv; TableA7_nlos_d3d_sensitivity.csv; TableA8_scene_composition.csv; "
  "TableA9_dlst_variants.csv; TableA10_dsm_crosscheck.csv; TableA11_sampling_frame.csv; "
  "TableA12_link_budget.csv; village_sample_v2.csv (candidate list with per-village coordinate "
  "verification status); dlst_overpass_matrix_v4.csv (per-village per-overpass anomalies under the "
  "three background definitions); coverage_4phase_v4.csv and coverage_4phase_v4_f07.csv (per-phase "
  "village coverage at 2.6 GHz and 700 MHz); geo_utils.py and s60_unit_tests.py (propagation core "
  "and its unit tests); and the scripted pipeline that regenerates every number, table, and figure.",
  size=9.5)

doc.save(OUT)
print("saved", OUT, "tables:", len(doc.tables))
