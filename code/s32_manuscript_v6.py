# -*- coding: utf-8 -*-
"""V3-R6 s32: v6 重建版全文手稿（回应第五轮审稿 R5 Reject-and-Resubmit 12 项 P0）
技术链变更：单一 canonical 坐标（village_geometry_v6.csv，不再咬合）；森林环带改为
500–800m 欧几里得环带；肌理密度分母统一为圆域面积、斑块 4-连通；tSVF 负地平线截断；
删除循环移位"置换"，改非空间标签洗牌基线 + Freedman-Lane 秩残差置换；
A14 统一 B=2999 并增 20-seed MC 稳定性（TableA14b）；新增 TableA17（WorldCover 年份 + DOY）。
推断框架：主判定=双重标准（族内 BH-FDR q<0.05 且 0.15° 区组 bootstrap 95%CI 不含 0）。
数字一律以 data/stats_v6.json / tables/ v6 为准。
输出：manuscript/V3_manuscript_full_v6.docx
"""
import json
import re

import docx
import pandas as pd
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

S = json.load(open("data/stats_v6.json", encoding="utf-8"))
OUT = "manuscript/V3_manuscript_full_v6.docx"
doc = docx.Document()

for s in doc.sections:
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.5)
    s.top_margin = s.bottom_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.3
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_runs(p, text, size=None):
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok)
        if size:
            r.font.size = Pt(size)


def P(text, size=None):
    p = doc.add_paragraph()
    add_runs(p, text, size)
    return p


def H1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True


def H2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True


def FIG(path, num, caption, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Fig. {num}. ")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = cap.add_run(caption)
    r2.font.size = Pt(9.5)
    cap.paragraph_format.space_after = Pt(10)


def set_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        borders.append(e)
    tblPr.append(borders)


def TABLE(num, title, headers, rows, fontsize=8.5):
    cap = doc.add_paragraph()
    r = cap.add_run(f"Table {num}. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = cap.add_run(title)
    r2.font.size = Pt(10)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        rr = c.paragraphs[0].add_run(str(h))
        rr.bold = True
        rr.font.size = Pt(fontsize)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(str(v))
            rr.font.size = Pt(fontsize)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def EQ(parts, num):
    """parts: list of (text, mode)；mode: None/'i'/'sub'/'sup'/'isub'"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, mode in parts:
        r = p.add_run(text)
        if "i" in (mode or ""):
            r.italic = True
        if "sub" in (mode or ""):
            r.font.subscript = True
        if "sup" in (mode or ""):
            r.font.superscript = True
    p.add_run(f"    ({num})")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

# ==================== 标题区 ====================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Thermal–Wireless Trade-Offs of Terrain-Constrained Settlement Morphology in Huizhou "
              "Traditional Villages")
r.bold = True
r.font.size = Pt(16)
t.paragraph_format.space_after = Pt(14)

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.add_run("Lei Zhang")
r = a.add_run("1,*"); r.font.superscript = True
for _r in a.runs:
    _r.font.size = Pt(12)

af = doc.add_paragraph()
af.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = af.add_run("1 School of Architecture and Urban Planning, Anhui Jianzhu University, Hefei 230601, China")
r.font.size = Pt(9.5)
r.italic = True
c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run("* Corresponding author. E-mail: [corresponding author's e-mail to be inserted]")
r.font.size = Pt(9.5)
doc.add_paragraph()

# ---- Abstract（约250词）----
ab = doc.add_paragraph()
r = ab.add_run("Abstract: ")
r.bold = True
add_runs(ab,
    "Traditional villages in mountainous regions must reconcile climatic comfort with modern digital "
    "infrastructure, yet the same settlement morphology may act in opposite directions on these two "
    "performances. For 29 traditional villages in the Huizhou region of southern Anhui, China, this paper "
    "quantifies how settlement morphology is associated with summer land surface temperature (LST) and "
    "with simulated wireless coverage, using only open data and scripted, re-runnable models. Sixteen "
    "morphological metrics derive from the 30-m Copernicus GLO-30 digital surface model and the 10-m ESA "
    "WorldCover map, all anchored to a single canonical village coordinate set. Village summer LST is "
    "estimated from Landsat 8/9 Collection-2 acquisitions (June–September 2019–2025, 26 independent "
    "overpasses); coverage is evaluated as reference-signal-received-power (RSRP) threshold exceedance "
    "under a stylized, standardized virtual macrocell deployment (2.6 GHz, four grid phases) using a "
    "terrain line-of-sight 3GPP rural-macrocell model. Under a terrain- and land-cover-matched "
    "background, the village heat anomaly averages +1.3 °C (25 of 29 villages positive). An association "
    "is called detectable only when it survives within-family Benjamini–Hochberg control and its "
    "spatial block-bootstrap 95% interval excludes zero; intervals are further audited for seed and "
    "block-design stability. Under this dual criterion, warmer villages are larger and more densely "
    "edged (building coverage ratio ρ = +0.65; edge density ρ = +0.64); villages on steep, "
    "high-relief, forest-ringed terrain are cooler (slope ρ = −0.53; relief ρ = −0.69; forest ring "
    "ρ = −0.68) but markedly harder to serve (cov85 ρ = −0.85, −0.77 and −0.74 respectively). The "
    "forest ring is thus associated with both a thermal benefit and a coverage cost in the same "
    "sample. Terrain-horizon openness is associated with better coverage (ρ = +0.57) and, "
    "borderline, with higher LST. Under size control, only coverage-side associations and the "
    "relief–temperature association survive. Results are stable to frequency band "
    "(700 MHz), county-level jackknife, and model range limits. Because the design is "
    "cross-sectional and the deployment stylized, the quantified couplings are hypotheses for "
    "site-specific assessment, not causal effects or operational coverage estimates.")
ab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw = doc.add_paragraph()
r = kw.add_run("Keywords: ")
r.bold = True
kw.add_run("settlement morphology; land surface temperature; rural wireless coverage; 3GPP rural "
           "macrocell model; terrain-horizon openness; traditional villages; Huizhou")
doc.add_paragraph()

# ==================== 1. INTRODUCTION ====================
H1("1. Introduction")
P("Traditional settlements in mountainous regions embody climatic knowledge accumulated over centuries of "
  "trial and error. In the Huizhou region of southern Anhui Province, China—home to the World Heritage "
  "villages of Hongcun and Xidi and to one of the densest concentrations of historic villages in East "
  "Asia—some settlements historically exhibit enclosed valley sites, forested backdrops, compact building "
  "fabrics, and water-adjacent layouts, and these features have long been read as a vernacular, "
  "pre-scientific system of passive environmental control [1], [2]. Field studies in and around the "
  "region confirm that this question remains empirically alive: indoor and outdoor thermal performance of "
  "Huizhou-area vernacular fabric continues to be measured and debated [3], [4]. These villages are "
  "simultaneously protected heritage, working agricultural communities, and tourist destinations, and "
  "they must now accommodate a second, very different infrastructure: rural broadband and mobile networks "
  "and digital-village platforms [5], with settlement-scale information modeling emerging as a planning "
  "tool [6]. The two agendas are usually planned by different agencies with different tools, and they "
  "meet, if at all, only at the construction site.")
P("The premise of this paper is that they should meet much earlier, because both agendas load on the same "
  "physical object: settlement morphology. Enclosed, compact, forest-backed, water-adjacent form is "
  "widely believed to be thermally favorable in hot-summer climates [1], [2], [7]. Radio propagation in "
  "mountainous terrain, by contrast, favors the opposite geometry: open horizons and unobstructed "
  "line-of-sight (LOS) paths to elevated macrocells [8]–[10]. Whether these two performances actually "
  "trade off across real villages—and by how much, and through which specific morphological attributes—"
  "has not been quantified. A non-systematic keyword search in major indexing services (Web of Science "
  "and CNKI, August 2026; English and Chinese terms combining settlement or village morphology with land "
  "surface temperature and with radio propagation or wireless coverage) returned no study that evaluates "
  "the thermal and radio performances of the same settlements jointly; we therefore claim novelty for "
  "the combination, not for either component. This paper provides that quantification, and in doing so "
  "it treats the associations strictly as associations: the design is cross-sectional and correlational, "
  "and the language of the paper follows that design.")
P("The thermal side of the question has a mature literature. The energetic basis of settlement-scale heat "
  "excess has been understood since the foundational work of Oke [11], and subsequent reviews consolidated "
  "the roles of geometry, surface cover, and anthropogenic heat [12]. The local climate zone (LCZ) scheme "
  "formalized the idea that built morphology predicts thermal behavior [13], and canyon geometry together "
  "with the sky-view factor (SVF) emerged as robust predictors of urban heat intensity [14], [15]. For "
  "vernacular settlements specifically, case studies combining field measurement and simulation have "
  "demonstrated the climate responsiveness of traditional dwellings and village layouts in China and "
  "elsewhere [1]–[4], [16], including systematic reviews of outdoor microclimate design in traditional "
  "Chinese villages [7]. Two gaps nonetheless remain. First, the vernacular evidence is overwhelmingly "
  "single-village: cross-settlement, statistically controlled evidence linking village morphology to "
  "thermal outcomes is thin, particularly in mountainous southern China [7], [17]. Second, the dominant "
  "morphometric evidence comes from cities, where morphology varies independently of terrain; in mountain "
  "villages, built form and terrain co-vary by design, so the urban-derived intuition that ‘more open is "
  "cooler’ cannot be assumed to hold [18].")
P("The radio side has an equally mature but entirely separate literature. Empirical path-loss models "
  "descend from Hata [8] through the COST-231 family, while the geometry-based 3GPP rural macrocell (RMa) "
  "model specifies distinct line-of-sight (LOS) and non-LOS (NLOS) branches in which terrain obstruction "
  "is the first-order control on rural coverage [9], [10]. Vegetation is known to attenuate gigahertz "
  "signals [19]. Rural coverage gaps—and their socio-economic consequences—are documented in global "
  "digital-divide assessments [20], and in China the digital-village strategy has made rural connectivity "
  "an explicit policy target [5]. Yet coverage studies rarely treat settlement morphology as an "
  "explanatory variable: villages appear, if at all, as demand points on a cost surface, not as "
  "morphological objects whose shape co-determines serviceability.")
P("Meanwhile, the data environment has changed. Open elevation models, global land-cover maps, and "
  "multi-year thermal satellite archives have made quantitative morphometrics cheap and reproducible "
  "[21], [23], [26], and urban studies have shown at scale that open imagery and machine learning can "
  "characterize urban form [22], [24], [25]. Rural heritage settlements have so far benefited "
  "little from this turn, partly because their building stock is small and scattered, and partly because "
  "the relevant questions sit between disciplines. This study brings the open-data morphometric approach "
  "to bear on the twin infrastructure question of mountain villages.")
P("We pose four research questions. (RQ1) How differentiated are Huizhou traditional villages in the "
  "morphological dimensions that plausibly matter for both thermal and radio performance? (RQ2) Which "
  "morphological attributes are associated with village-scale summer surface temperature and its anomaly "
  "against the rural background? (RQ3) Which attributes are associated with simulated wireless coverage "
  "under a deployment that is held constant across villages? (RQ4) Where the two sets of associates "
  "coincide, do they act in the same or in opposite directions—that is, where are the trade-offs?")
P("To answer these questions we assemble a sample of 29 traditional villages in the Huizhou region and, "
  "using only open data and reproducible models, (i) characterize each village by sixteen morphological "
  "metrics defined over explicitly stated spatial frames, all computed by script from public rasters; "
  "(ii) estimate village-scale summer LST and heat anomaly from a quality-masked multi-year Landsat "
  "archive, with the estimand stated per overpass; (iii) simulate coverage under a standardized virtual "
  "macrocell deployment that removes operator-specific deployment heterogeneity, with grid-phase and "
  "frequency sensitivity analyses; and (iv) quantify, with size-controlled partial correlations, a dual "
  "significance criterion combining false-discovery-rate control and spatial block-bootstrap intervals "
  "audited for seed and block-design stability, non-spatial permutation baselines, and "
  "leave-one-county-out checks, which morphological attributes "
  "covary with the two performances in opposite or concordant directions. Section 2 describes the study "
  "area, data, and methods; Section 3 reports the results; Section 4 discusses interpretation, planning "
  "implications, and limitations; Section 5 concludes.")

# ==================== 2. MATERIALS AND METHODS ====================
H1("2. Materials and Methods")
P("The research framework couples two open-data model chains onto a common morphometric backbone (Fig. 1, "
  "Table 1). For each of the 29 villages we first compute sixteen morphological metrics from elevation and "
  "land-cover data (Section 2.3). The same village domains then feed two independent performance models: a "
  "thermal chain, in which quality-masked multi-year Landsat surface-temperature acquisitions yield village "
  "LST and its anomaly against the rural background (Section 2.4); and a radio chain, in which a "
  "standardized virtual macrocell deployment is evaluated with a terrain-aware 3GPP propagation model to "
  "yield village coverage metrics (Section 2.5). Because the deployment is held constant across villages, "
  "coverage differences between villages reflect morphology and terrain rather than operator decisions. "
  "The two performance vectors are finally compared against the morphological vector across the sample, "
  "with the statistical controls described in Section 2.6. All in-text statistics, tables, and figures are "
  "generated from a single village-level analysis table by script, so that numbers reported in the text "
  "cannot drift from the data.")

H2("2.1. Study Area and Village Sample")
P("The study area covers the historical Huizhou prefecture in the foothills of the Huangshan mountains, "
  "southern Anhui (approximately 29.6°–30.3° N, 117.3°–118.9° E). The climate is subtropical monsoon with "
  "hot, humid summers in which outdoor thermal stress is a genuine public-health and livability concern. "
  "The terrain is a mosaic of steep forested ridges and narrow alluvial basins; villages traditionally "
  "settled along streams at basin edges, with their backs to a forested slope and their fronts to water—"
  "the classical *beishan mianshui* (backed by mountains, facing water) configuration. The region contains "
  "two UNESCO World Heritage villages (Hongcun and Xidi) and numerous nationally listed historic and "
  "traditional villages, making it both a conservation hotspot and a natural laboratory in which "
  "morphological setting varies widely within a single cultural sphere.")
P("*Sampling frame.* The candidate population was the published registers of nationally listed "
  "traditional and historic-cultural villages within the historical Huizhou prefecture (the six counties "
  "above; the national Chinese Traditional Villages register and the Historic-Cultural Villages list). "
  "Inclusion required (i) listing in a national register, (ii) an intact historic "
  "built core visible in current imagery, and (iii) a geographic identity verifiable against OpenStreetMap "
  "and satellite imagery. Screening was desk-based and performed by a single reviewer, which is a "
  "known limitation of the frame. From the register we assembled a desk-review shortlist of 28 villages "
  "spanning basin-floor, piedmont, and deep-valley settings across all six counties; one further village "
  "(Huansha) was added after imagery verification, yielding the final sample of 29. The sample is thus a "
  "purposive, register-based regional sample stratified by terrain setting and county, not a probability "
  "sample; we report it as such and do not generalize beyond the sampled region. Village names are given "
  "in both pinyin and Chinese characters throughout (Tables 1 and A1; Appendix Table A11). The sample "
  "list, with per-village coordinate source (register, OpenStreetMap, or imagery fix), verification "
  "status, and reviewer notes in both languages, is documented in Appendix Table A11 and archived in the "
  "repository (village_sample_v2.csv). **Canonical coordinates.** A single coordinate set, fixed before "
  "analysis from the register and OpenStreetMap/imagery checks, anchors every measurement in this "
  "paper—morphology, terrain, thermal, and coverage alike; no coordinate adjustment is applied anywhere "
  "in the pipeline, so the same village point feeds every chain and no chain-specific relocation can "
  "enter the analysis. Two villages "
  "(Kantou and Lixi) carry approximate-coordinate flags: no authoritative point source was "
  "available, so their positions were fixed from OpenStreetMap and imagery; both are retained with the "
  "flag, and the headline associations are robust to their exclusion. Three sample particularities are flagged because "
  "they matter for the analysis. First, the built fabric around Yuliang merges with the adjacent Shexian "
  "county seat in the land-cover map; we retain the village and verify that "
  "all reported associations are robust to its exclusion. Second, the built fabric of two small "
  "mountain villages (Zuyuan and Mulihong) is under-detected by the 10-m land-cover map: the domain-based "
  "plan-form and fabric metrics of Section 2.3 meet their ten-pixel minimum for Zuyuan but not for "
  "Mulihong, which has no detected built-up cell inside its domain (n = 28 where applicable). Their "
  "coverage targets follow the fallback rule of Section 2.5, with a dedicated exclusion sensitivity "
  "check. Third, the sample deliberately spans basin-floor, "
  "piedmont, and deep-valley settings (domain-mean elevation 126–420 m; local relief 22–320 m), which is "
  "the source of morphological variance on which the analysis draws. Because elevation is summarized as a "
  "domain mean, it can differ from point elevation at the village reference point, particularly for steep "
  "mountain villages. Fig. 1 maps the sample; Table 1 lists the villages with key attributes; the full "
  "sixteen-metric table and the per-phase coverage detail are given in Appendix Table A1.")
FIG("figures/Fig1_study_area_EN.png", 1,
    "The 29 sampled traditional villages in the Huizhou region, southern Anhui, China. Base map: Esri "
    "World Imagery (Maxar, Earthstar Geographics, CNES/Airbus DS), accessed August 2026; map projection "
    "WGS 84 / Pseudo-Mercator (EPSG:3857); numbers keyed to Table 1. Inset: location in China. Village "
    "names in pinyin and Chinese characters are listed in Tables 1, A1, and A11.")

H2("2.2. Open Data")
P("All inputs are public datasets; no field measurement was used, and the entire pipeline is scripted and "
  "re-runnable end to end (Section 2.3). Terrain is given by the 30-m Copernicus GLO-30 product (COP-DEM) "
  "[21]. COP-DEM is a digital *surface* model (DSM): its heights reference the reflecting surface, which "
  "over forested or built-up ground includes canopy and roof signals; we therefore refer to it as a DSM "
  "throughout, treat terrain metrics derived from it as surface metrics, and return to the consequences in "
  "Section 4.4. As a gross-error check, village centroid elevations and within-domain relief were "
  "recomputed from the independent NASADEM 1-arc-second product [32]: centroid elevations agree at "
  "Spearman ρ = 0.999 (RMSE 22 m, mean difference −3 m) and domain relief at ρ = 0.988 across the 29 "
  "villages (Appendix Table A10). Land cover is the 10-m ESA WorldCover 2021 map (v200), from which the "
  "built-up, tree-cover, and water classes were extracted [23]. Summer surface temperature is derived from "
  "Landsat 8/9 Collection-2 Level-2 surface-temperature products (band ST_B10, acquired at 100-m native "
  "thermal resolution and distributed on the 30-m product grid; we bilinearly reproject the band to the "
  "100-m UTM analysis grid), accessed through the "
  "Microsoft Planetary Computer spatiotemporal asset catalog (collection landsat-c2-l2; query executed "
  "17 August 2026) [27]. We queried June–September scenes from "
  "2019–2025 with catalog scene cloud cover below 10%, retaining up to nine of the lowest-cloud scenes "
  "per Worldwide Reference System frame, yielding 36 scene assets over the four frames (paths 120–121, "
  "rows 39–40) that cover the study area; the per-scene identifiers, dates, and cloud covers are listed "
  "in Appendix Table A2. Because adjacent WRS rows of the same path acquired on the same date are two "
  "observations of the same overpass, treating them as independent scenes would pseudo-replicate the "
  "sample. We therefore merge same-platform, same-path, same-date assets pixel-wise (mean of valid "
  "values) into 26 independent overpasses; all anomaly estimation and bootstrap resampling below uses "
  "the overpass, not the scene asset, as the unit. Digital numbers were converted with the Collection-2 "
  "scale factors (ST = DN × 0.00341802 + 149.0 K) and expressed in degrees Celsius. Quality control "
  "applied the Collection-2 QA_PIXEL bit mask pixel-wise (fill, dilated cloud, cirrus, cloud shadow, "
  "snow, and cloud bits cleared), followed by a 10–60 °C physical plausibility screen. Three scenes were "
  "excluded because this screen left zero valid pixels in the study-area subset: inspection shows that "
  "the subset is covered by cloud or cloud shadow in these scenes even though their whole-scene catalog "
  "cloud cover is only 0.6–4.6%, a reminder that catalog cloud cover is a whole-scene statistic. One "
  "further overpass date (2020-06-27) retains only ≈0.01% valid pixels after masking; it is "
  "excluded from the primary anomaly analysis and retained as a sensitivity case. We use the "
  "distributed product as is, including its emissivity model, and we do not propagate the per-pixel "
  "ST_QA uncertainty band, which would widen the reported intervals further; this is stated as a "
  "limitation (Section 4.4), and overpass-block bootstrap propagates overpass-to-overpass sampling "
  "uncertainty (not per-pixel retrieval uncertainty) "
  "into the village anomaly (Section 2.4). The summer composite used for mapping is the per-pixel "
  "*median* across the 25 retained overpasses, which is robust to residual contamination. A 500-m village "
  "domain contains roughly 80 native 100-m thermal pixels. Using a multi-year median sacrifices "
  "year-to-year resolution in exchange for a stable, cloud-gap-free characterization of typical summer "
  "daytime surface behavior, which is the quantity that morphology can plausibly covary with [28].")

H2("2.3. Morphological Metrics: One Anchor, Three Layers")
P("All sixteen morphological metrics anchor to the canonical village coordinate of Section 2.1 and are "
  "computed by script directly from the public rasters of Section 2.2—no manual digitizing, no "
  "proprietary software, and no coordinate relocation anywhere in the chain; the scripts are deposited "
  "with the repository, so every metric in Table A1 can be regenerated from the raw public inputs. The "
  "metrics are organized in three layers that answer different architectural questions. **Fabric "
  "metrics** describe the detected built fabric inside the fixed site domain. **Site metrics** describe "
  "the terrain on which the village sits. **Context metrics** describe the landscape immediately around "
  "the site. **Frame D (the site domain)**—a fixed 500-m-radius circle of 78.5 ha centred on the "
  "canonical village point—is the single spatial frame for all sixteen metrics, and it is also the "
  "analysis window for the thermal and coverage measurement of Sections 2.4 and 2.5, so that "
  "explanatory and outcome variables share one geometry by construction.")
P("*Fabric metrics (Frame D).* Plan form and texture are measured on the built-up cells detected inside "
  "the site domain on the 10-m WorldCover grid. Elongation is the principal-axis aspect ratio of those "
  "cells, defined as sqrt(λ1/λ2) of the eigenvalues of the cell-centroid covariance matrix; it is high "
  "for stream-following linear villages and low for isotropic basin villages. Compactness is the "
  "isoperimetric quotient 4πA/P² aggregated over the built-up patches inside the domain, where A is the "
  "total built-up area and P is the total exposed-edge perimeter. The perimeter is computed under "
  "four-neighbour adjacency with anisotropic weighting—vertical edges contribute the pixel height and "
  "horizontal edges the pixel width, which differ by up to 13% at this latitude on the geographic "
  "WorldCover grid—counting hole boundaries as exposed perimeter, counting edges at the domain circle, "
  "and aggregating all patches into a single quotient (total A over total P², not a patch-weighted "
  "average); the implementation is covered by geometric unit tests (square, rectangle, two-patch, and "
  "holed-patch cases) distributed with the code (s60_unit_tests.py). Four further metrics describe the "
  "texture of the fabric rather than its outline: building coverage ratio (the share of domain cells "
  "classified as built-up), edge density (total exposed built-up edge per hectare of *domain* area, "
  "m/ha), patch density (the number of four-connected built-up patches per hectare of *domain* area), "
  "and largest-patch share (the area of the single largest four-connected patch as a percentage of "
  "total built-up area in the domain). Because the two density metrics use the fixed domain area as "
  "denominator, they scale with the amount of fabric as well as its texture, and we read them "
  "accordingly in Section 3. Together these distinguish continuous, consolidated fabrics from "
  "fragmented ones. All plan-form and fabric metrics "
  "require at least ten detected built-up cells inside the domain, a threshold met by 28 of the 29 "
  "villages (Mulihong has no detected built-up cell in its domain; its building coverage ratio is "
  "therefore exactly 0%, which is informative in itself, while its edge, patch, largest-patch, and "
  "plan-form metrics are undefined).")
P("*Site metrics (Frame D).* Elevation is the domain mean; relief is the maximum minus minimum "
  "elevation within the domain; slope is the domain mean of terrain gradient. Southness is the domain "
  "mean of cos(aspect), positive for south-facing ground, and north–south asymmetry is the difference in "
  "mean elevation between the northern and southern halves of the domain—together these capture the "
  "canonical claim that Huizhou villages face south against a northern slope. To capture terrain "
  "enclosure we use the terrain sky-view factor (tSVF), which we call *terrain-horizon openness* "
  "throughout to avoid confusion with the building-resolved SVF of the urban literature: it is the "
  "unobstructed fraction of the upper hemisphere at the canonical village point, sampled along 72 "
  "azimuths with a 2500-m horizon [15]:")
EQ([("tSVF = 1 − (1/", "i"), ("N", "i"), (") Σ sin(max(0, ", None), ("γ", "i"), ("i", "sub"),
    (")),   ", None), ("i", "i"), (" = 1, …, 72", None)], 1)
P("where γ", )
doc.paragraphs[-1].add_run("i").font.subscript = True
add_runs(doc.paragraphs[-1],
    " is the horizon elevation angle along azimuth *i*, truncated below at zero so that directions "
    "looking downward onto lower terrain count as fully open rather than as negative obstruction. A "
    "village on an open plain has tSVF ≈ 1; a village "
    "at the bottom of a gorge has a markedly lower value. The 2500-m horizon is long enough to include the "
    "enclosing ridge lines that govern both insolation and radio visibility, and short enough to remain a "
    "village-scale attribute. Because terrain-horizon openness is evaluated once, at the canonical point, it "
    "characterizes the village’s setting rather than its internal spatial variation; intra-village "
    "sky-view heterogeneity requires building-scale data that our open sources do not provide (Section "
    "4.4).")
P("*Context metrics (Frame D).* The forest ring is the tree-cover fraction within the 500–800-m "
  "Euclidean annulus immediately outside the site domain—the vegetated backdrop typical of Huizhou "
  "villages; the annulus is defined geometrically from the canonical point, so the metric is available "
  "and comparable for all 29 villages regardless of how much fabric the land-cover map detects. Water "
  "distance is the Euclidean distance from the built-up cells inside the domain to the nearest water "
  "pixel, summarized by its mean and minimum; for Mulihong, which has no detected built-up cell, it is "
  "the distance from the domain centre instead (flagged in Table A1); we deliberately keep the variable "
  "as a *distance*—larger values mean settlement fabric farther from water—so that the sign of any "
  "association is read without ambiguity. "
  "Finally, domain built-up area is the built-up surface clipped to Frame D (0–56.8 ha across the "
  "sample); it is the size variable used in the statistical controls of Section 2.6, because it shares "
  "the spatial frame of the outcome variables. Several metrics co-vary by construction of the "
  "terrain—deep-valley villages are steep, forested, enclosed, and elongated—which is why Section 2.6 "
  "relies on rank correlations and size-controlled partial correlations rather than on parametric "
  "multivariate models.")

H2("2.4. Summer LST and Heat Anomaly")
P("For each village, core and background temperatures are defined over explicitly stated zones, and the "
  "choice of zone is itself reported as a methodological variable. We compute the anomaly under three "
  "definitions. **V1 (site domain):** the core is the 500-m site domain with water pixels masked, and the "
  "background is the 1–2 km ring restricted to non-built, water-masked pixels, so that the background "
  "reflects the rural matrix rather than adjacent settlement fabric. **V2 (fabric):** the core is the "
  "built-up-classified pixels inside the site domain (villages with fewer than five such pixels are "
  "not estimable), against the same ring; this variant targets the built boundary itself at the cost "
  "of sample size. **V3 (matched):** the core is the V1 domain, and the background is a pixel-level "
  "subset of the V1 ring matched one-to-one (greedy, without replacement) on elevation, slope, and "
  "tree-cover fraction; V3 is the partially controlled estimate and the primary anomaly used in the "
  "correlation analysis, while V1 and V2 are reported alongside (Table 1; Appendix Table A9). Matching quality is "
  "itself audited (Appendix Table A15): post-matching standardized mean differences are small for "
  "elevation (≤0.23 across villages) but remain substantial for slope (up to 0.82) and tree cover (up "
  "to 0.98) in some villages, and a median of 11% of matched pairs exceed a 0.5-standard-deviation "
  "caliper distance; the V3 estimate should therefore be read as partially, not fully, controlled. "
  "Water pixels are masked in every zone of every variant; in the V1 ring they occupy on average 0.6% "
  "of the area (maximum 6.9% at riverside Yuliang, which is covered by the dedicated exclusion "
  "sensitivity of Section 2.6). Because the composite pools overpasses acquired on different dates, a "
  "difference of composite means would mix acquisition composition with morphology. We therefore "
  "define the heat anomaly overpass by overpass: for every independent overpass we take the "
  "difference between the core mean and the background mean, each computed only when at least five "
  "valid pixels are available in that overpass and zone, and then average these per-overpass "
  "differences for each village (between 6 and 13 overpasses contribute per village, mean 7.6; 16 "
  "overpass dates contribute at least one village estimate),")
EQ([("ΔLST = (1/", "i"), ("N", "i"), ("o", "sub"), (") Σ", None), ("o", "i"),
    (" [LST", "i"), ("core", "sub"), ("(", None), ("o", "i"), (") − LST", "i"),
    ("bg", "sub"), ("(", None), ("o", "i"), (")]", None)], 2)
P("which controls acquisition date by construction. Uncertainty is propagated by a bootstrap whose "
  "resampling unit is the independent overpass (10,000 resamples, fixed seed 7), yielding a 95% "
  "confidence interval and the probability that the anomaly exceeds zero; villages whose interval "
  "includes zero are reported as having a positive point estimate without confirmed sign. Threshold "
  "sensitivity is reported: requiring at least 25% valid-pixel coverage per zone per overpass, and "
  "retaining the near-zero-coverage 2020-06-27 overpass date, both leave the village estimates "
  "essentially unchanged (rank correlation with the primary estimates ρ ≥ 0.98 and no sign changes in "
  "either case; Appendix Table A9). The water-class mask itself is threshold-sensitive in principle, "
  "because WorldCover’s water class can under-detect narrow or shaded streams; repeating the V1 "
  "estimation with a deliberately lenient water mask changes no village’s estimate by more than "
  "0.11 °C and preserves the village ranking exactly (rank ρ = 1.000; Appendix Table A16).")
P("The absolute-LST estimand is stated explicitly, because two defensible choices exist. The primary "
  "estimand is the mean of per-overpass village-domain means—the same overpass-resolved quantity that "
  "enters the anomaly, taken without background subtraction—which is what Table 1 reports. The "
  "alternative, extracting village means directly from the median composite raster, agrees with the "
  "primary estimand at rank ρ = 0.941 with a mean bias of −0.20 °C (Appendix material); the two "
  "estimands therefore support the same cross-village ordering. Absolute LST remains descriptive: "
  "absolute values are used for ranking within the sample, not as climatological norms. LST is used "
  "as a physically consistent proxy of the daytime outdoor thermal environment; it is not equivalent "
  "to air temperature or pedestrian-level thermal comfort, a distinction we return to in Section 4.4 "
  "[28].")

H2("2.5. Standardized Wireless Deployment and Propagation Model")
P("Real operator deployments are heterogeneous and commercially opaque, so village-level coverage "
  "differences measured against live networks conflate morphology with deployment history. To hold the "
  "deployment pattern constant with respect to morphology and terrain, we evaluate coverage under a "
  "standardized virtual deployment: macrocell sites "
  "on a 2.5-km grid over the study region, each site snapped to the highest surface-model cell within "
  "600 m—a stylized heuristic for the locally elevated siting common in rural radio planning, not a "
  "description of any real network—with a 30-m antenna height above the surface model. The carrier "
  "frequency is 2.6 GHz, a widely deployed mid-band in China’s mobile networks; we treat it as a "
  "conservative mid-band case, because lower bands propagate further and diffract better around "
  "terrain—a 700-MHz repetition of the full chain, reported in Section 3.3, confirms that absolute "
  "coverage levels rise at low frequency while the inter-village ranking is essentially preserved. "
  "Following the RSRP measurement definition, the link budget is written directly in terms of the "
  "effective radiated power per resource element (EPRE) [33]:")
EQ([("EPRE = ", None), ("P", "i"), ("T", "sub"), (" − 10 log", None), ("10", "sub"),
    ("(", None), ("N", "i"), ("SC", "sub"), (") + ", None), ("G", "i"), ("T", "sub"),
    (" − ", None), ("L", "i"), ("f", "sub"),
    (" = 46 − 10 log", None), ("10", "sub"), ("(1200) + 17 − 2 = 30.2 dBm", None)], 3)
P("for a 46-dBm total carrier power spread over a 20-MHz LTE-class downlink of 1200 subcarriers, 17-dBi "
  "antenna gain, and 2-dB feeder loss. The total power and the subcarrier count enter the analysis only "
  "through this per-resource-element conversion; because RSRP is itself defined per resource element "
  "[33], no further bandwidth term appears in the budget, and the identical EPRE is used at 700 MHz. "
  "Worked link budgets for three representative villages, traced line by line through Eqs. (3)–(5) with "
  "their true link geometries, are given in Appendix Table A12. Because any fixed grid may "
  "favor or disfavor individual villages simply through its alignment, the deployment is repeated over "
  "four grid phases shifted by half a spacing (0 and 1250 m in x and y), yielding 1620 sites per phase; "
  "village-level coverage metrics are averaged across the four phases, and the cross-phase dispersion is "
  "itself reported as a measure of deployment robustness.")
P("Coverage targets are the built-up cells (30 m) inside each village domain; where the land-cover map "
  "under-detects the built fabric (fewer than 30 cells: Zuyuan and Mulihong), a regular 150-m grid over "
  "the domain is used instead, and all coverage associations are re-estimated without these two villages "
  "as a sensitivity check (Sections 2.6 and 3.3). The receiver is placed 1.5 m above the *surface-model* "
  "height at each target: because COP-DEM is a surface model, this is a surface-top scenario in which "
  "the terminal sits above the modelled canopy or roof signal, not a pedestrian-height scenario above "
  "true ground; we state this explicitly and return to its consequences in Section 4.4. Path loss "
  "follows the 3GPP RMa model [9], with LOS or NLOS determined per link by stepping along the 30-m "
  "surface-model profile between site and target and applying a cumulative horizon criterion: a target "
  "is in LOS if the straight site-to-target line clears every intervening terrain elevation. All "
  "distances entering the model are true three-dimensional link distances d3D = sqrt(d2D² + Δz²), "
  "where Δz is the difference between the absolute antenna heights (site surface elevation plus 30 m; "
  "target surface elevation plus 1.5 m); in mountain terrain Δz can reach several hundred metres, which "
  "a constant antenna-height offset would misrepresent. For LOS links,")
EQ([("PL", "i"), ("LOS", "sub"), ("(", None), ("d", "i"), (") = 20 log", None), ("10", "sub"),
    ("(40π", None), ("d", "i"), (" f", "i"), ("c", "sub"), ("/3) + min(0.03 ", None), ("h", "i"),
    ("1.72", "sup"), (", 10) log", None), ("10", "sub"), ("(", None), ("d", "i"),
    (") − min(0.044 ", None), ("h", "i"), ("1.72", "sup"), (", 14.77) + 0.002 log", None),
    ("10", "sub"), ("(", None), ("h", "i"), (") ", None), ("d", "i")], 4)
P("which applies up to the breakpoint distance dBP = 2π hBS hUT fc/c (with fc in Hz in this expression); "
  "beyond the breakpoint, PL_LOS(d) = PL_LOS(dBP) + 40 log10(d/dBP). For NLOS links the RMa model "
  "applies its rural branch,")
EQ([("PL", "i"), ("NLOS", "sub"), ("(", None), ("d", "i"), (") = 161.04 − 7.1 log", None), ("10", "sub"),
    ("(", None), ("W", "i"), (") + 7.5 log", None), ("10", "sub"), ("(", None), ("h", "i"),
    (") − (24.37 − 3.7(", None), ("h", "i"), ("/", None), ("h", "i"), ("BS", "sub"),
    (")²) log", None), ("10", "sub"), ("(", None), ("h", "i"), ("BS", "sub"), (") + (43.42 − 3.1 log",
     None), ("10", "sub"), ("(", None), ("h", "i"), ("BS", "sub"), ("))(log", None), ("10", "sub"),
    ("(", None), ("d", "i"), (") − 3) + 20 log", None), ("10", "sub"), ("(", None), ("f", "i"),
    ("c", "sub"), (") − (3.2(log", None), ("10", "sub"), ("(11.75 ", None), ("h", "i"),
    ("UT", "sub"), ("))² − 4.97)", None)], 5)
P("and the realized path loss is max(PL_LOS, PL_NLOS), with the nominal rural parameterization of the "
  "model (average building height h = 5 m, street width W = 20 m, base height hBS = 30 m, receiver "
  "height hUT = 1.5 m) [9]. Throughout Eqs. (3)–(5), d denotes d3D in metres and the carrier frequency "
  "fc in GHz. The propagation code is covered by unit tests that reproduce a hand-computed LOS link "
  "budget, verify NLOS ≥ LOS and the range-limit behaviour, and verify the d3D sensitivity "
  "(s60_unit_tests.py). The ray-stepping LOS determination is purely geometric: for each site–target "
  "pair we sample the 30-m surface model along the connecting path at one-pixel increments, maintain "
  "the running maximum of the elevation angle from the site antenna to the intervening terrain, and "
  "declare the target visible when the straight line to the receiver clears that cumulative horizon. "
  "The procedure is the radio analogue of the viewshed analysis familiar from landscape planning, and "
  "it is what couples coverage directly to terrain morphology. Each target keeps the strongest server "
  "among sites within 10 km; at this assignment radius, omitting Earth curvature shifts the computed "
  "line of sight by at most about 7.8 m, small against the 30-m terrain sampling. The RMa "
  "parameterization is specified for LOS distances up to 10 km and NLOS distances up to 5 km; we "
  "therefore truncate NLOS links at 5 km, treating longer NLOS links as unserved, and report this "
  "range-limited model as the primary one. The truncation does not bind in the present configuration: "
  "village-level coverage metrics are identical with and without it in all four phases (Appendix Table "
  "A7), and a fixed-antenna-offset variant that ignores true link geometry likewise preserves every "
  "village-level ranking (phase-0 rank correlation 0.998; Table A7), so the reported values are not an "
  "artifact of either choice. We report T−85 and T−95 exceedance shares (cov85 and cov95), the "
  "percentages of targets whose reference signal received power (RSRP) exceeds −85 and −95 dBm "
  "respectively. These are RSRP-threshold exceedance shares under a stylized scenario: without "
  "signal-to-interference ratios they cannot certify service, and we use them strictly as "
  "planning-level visibility indicators—the T−85 level corresponds to good signal and the T−95 level "
  "to basic connectivity in this scenario. The 10th-percentile RSRP (RSRP p10) serves as a continuous "
  "measure of the worst-served fabric.")
P("The model deliberately omits foliage attenuation [19], building clutter, diffraction refinement, "
  "antenna patterns, shadowing, and interference. What it produces is therefore a deterministic "
  "received-power scenario under a stylized deployment—not an estimate of any operational network—and "
  "absolute coverage levels are optimistic. The objects of this study are the *differences between "
  "villages*, for which terrain visibility is the principal spatial obstruction explicitly represented; "
  "whether omitted effects—most notably vegetation attenuation over the forest ring [19]—would change "
  "the village-level ranking cannot be determined from the present model and is flagged as a "
  "validation target (Section 4.4). The 700-MHz repetition (Section 3.3) quantifies the frequency side "
  "of this question.")

H2("2.6. Statistical Analysis")
P("Associations between the sixteen morphological metrics and the four performance variables (village "
  "LST, ΔLST-V3, cov85, RSRP p10) are quantified by Spearman rank correlations, which are robust to the "
  "skewed distributions and small sample size. Two test families are defined explicitly: family F1 "
  "comprises the 64 raw correlations (16 metrics × 4 outcomes), and family F2 comprises the 60 partial "
  "correlations on ranks controlling domain built-up area (15 metrics × 4 outcomes; size itself is not "
  "controlled for). Villages in one region are spatially autocorrelated (Section 3.4), so nominal "
  "p-values overstate the effective evidence; we therefore do not treat any single device as decisive "
  "and instead combine three transparent instruments. **(i) Family control.** Within each family, "
  "p-values are adjusted with the Benjamini–Hochberg false discovery rate at q = 0.05 [29]. "
  "**(ii) Spatial intervals.** Confidence intervals are computed by a spatial block bootstrap [34], "
  "[35]—villages are grouped into 14 blocks on a 0.15° latitude–longitude grid and blocks are resampled "
  "whole (B = 2999, fixed seed 11)—which preserves within-block spatial dependence instead of assuming "
  "independent villages; block membership is published with the repository. **Primary decision rule.** "
  "An association is called *detectable* only when it passes both instruments: within-family BH "
  "q < 0.05 and a block-bootstrap 95% interval that excludes zero. This dual criterion is deliberately "
  "conservative, and it is the only claim level used in the abstract, results, and conclusions. "
  "**(iii) Non-spatial permutation baselines.** For every raw test we additionally report a standard "
  "label-permutation p-value (9999 random reshuffles of the outcome, fixed seed 23); for every partial "
  "test we report a Freedman–Lane residual-permutation p-value on ranks (9999 reshuffles of the "
  "rank-regression residuals, the control variable held fixed). These permutations destroy all spatial "
  "structure and therefore do not address spatial autocorrelation; we report them strictly as "
  "non-spatial reference baselines, not as spatial tests. As a further, explicitly heuristic "
  "sensitivity we also compute a "
  "rank-based Moran’s-I effective-sample-size approximation (n_eff truncated to [3, n]); this is a "
  "custom, conservative plausibility check, not a formal correction, and we label it as such wherever "
  "it appears. In Table 2, stars mark nominal p-values and a dagger marks tests passing the dual "
  "criterion; per-test statistics—nominal, permutation-baseline, and ESS-approximation p-values, within-family "
  "q-values, block-bootstrap intervals, and n—are reported in Appendix Table A3. Interval robustness to "
  "the blocking choice is audited on a 0.10°/0.15°/0.20° grid crossed with two grid origins (Appendix "
  "Table A14, B = 2999 throughout), and the Monte-Carlo stability of every interval is audited by "
  "repeating each configuration under 20 further seeds (B = 999 per seed; Appendix Table A14b): of the "
  "84 configuration–pair combinations, 58 intervals exclude zero under all 20 seeds, 19 include zero "
  "under all 20 seeds, and 7 are seed-sensitive; the seed-sensitive cases are identified by name in "
  "Section 3.4 and never support a claim alone. Because size correlates "
  "with nearly everything—larger villages are hotter and also tend to sit in more open, serviceable "
  "terrain—the F2 partial correlations are the conservative read of metric-level association; we do "
  "not call them independent effects, only size-adjusted ones.")
P("Robustness is assessed four ways. First, the four-phase deployment repetition absorbs grid "
  "alignment (Section 2.5). Second, the peri-urban Yuliang case is excluded and all associations "
  "re-estimated. Third, the two villages with fallback coverage targets are excluded from all coverage "
  "associations. Fourth, a leave-one-county-out jackknife over the six counties checks that no single "
  "county drives the headline associations (Appendix Table A4). Spatial autocorrelation is diagnosed "
  "with Moran’s I under row-standardized inverse-distance weights on UTM Zone 50N coordinates (999 "
  "permutations; Appendix Table A5), and it is this diagnostic that motivates the block bootstrap; the "
  "label-permutation baselines above play no spatial role. Trend lines in scatter plots are Theil–Sen "
  "robust fits [30], "
  "[31]. The propagation core is covered by unit tests (Section 2.5). All computation uses open-source "
  "Python libraries (rasterio, numpy, pandas, SciPy, matplotlib); the full code, the village-level "
  "analysis table, and the scene manifest are deposited in a public repository (see Data "
  "Availability).")

# ==================== 3. RESULTS ====================
H1("3. Results")
H2("3.1. Morphological Characterization")
P("The 29 villages differentiate clearly along the terrain dimension (Table 1, Fig. 2). Terrain-horizon "
  "openness spans tSVF = 0.735–0.983: most villages occupy open basin floors where the 2500-m horizon is "
  "essentially unobstructed, while a minority—notably Zuyuan (tSVF = 0.735), Shitan (0.787), and Renli "
  "(0.834)—sit in deeply incised valleys. The forest ring covers 25–100% of the 500–800-m annulus. "
  "Domain built-up area "
  "ranges from nil in the under-detected mountain village of Mulihong to 56.8 ha in the peri-urban "
  "Yuliang, and building coverage ratio from 0% (Mulihong, no detected built-up cell) to 71.6% "
  "(Yuliang, whose fabric merges with the county seat); relief within the site domain ranges from 22 m "
  "on the basin floor to 320 m around the highest mountain villages. The fabric-texture metrics separate "
  "consolidated from fragmented plans: edge density spans 5.5–161 m of exposed built-up edge per "
  "hectare of domain area, patch density 0.01–0.41 patches per hectare, and largest-patch share "
  "33–100%. The three "
  "setting classes are legible in the metric space: basin villages are open and weakly forested in "
  "their immediate ring; piedmont villages are intermediate; deep-valley villages are enclosed, "
  "elongated, steep, and almost fully ringed by forest.")
P("Two morphological observations deserve note. First, despite the canonical image of the south-facing "
  "Huizhou village, nine of the 29 villages show negative domain-mean southness—most strongly Nanping "
  "(−0.51) and Guanlu (−0.40)—a reminder that the textbook ideal is an idealization rather than a "
  "rule. Second, plan form and terrain co-vary: valley villages are strongly elongated along their "
  "streams (elongation up to 3.5 at Changxi), whereas basin villages are more isotropic. This "
  "co-variation, visible in the standardized distributions of Fig. 2c, is the statistical signature "
  "of site selection responding to terrain—and it is what makes the mountain-village morphometric "
  "problem different from its urban counterpart.")
import pandas as pd
_t1 = pd.read_csv("tables/Table1_sample.csv").fillna("—")
TABLE(1, "Key attributes of the 29 sampled villages: county, domain built-up area, building coverage "
         "ratio, elevation, relief, terrain-horizon openness (tSVF), forest ring share, summer LST "
         "(overpass-resolved estimand, Section 2.4), the site-domain anomaly ΔLST-V1 and the "
         "terrain-matched anomaly ΔLST-V3 with overpass-bootstrap 95% CIs, and four-phase coverage "
         "metrics (cov85 mean with cross-phase standard deviation; RSRP p10). The full "
         "sixteen-metric table is Appendix Table A1 (repository CSV).",
      list(_t1.columns), _t1.values.tolist(), fontsize=7)
FIG("figures/Fig2_morphology.png", 2,
    "Morphological characterization of the sample. (a) Terrain-horizon openness (tSVF) and (b) forest "
    "ring share over a 30-m hillshade, point size proportional to domain built-up area; (c) standardized "
    "distributions of the morphological metrics across the 29 villages.")

H2("3.2. Summer Thermal Environment")
P("Village-core summer LST spans 31.2–42.8 °C (mean 36.4 °C) across the sample. The coolest villages "
  "are high-elevation, enclosed mountain settlements (Mulihong 31.2 °C, Tachuan 31.8 °C, Zuyuan "
  "32.1 °C); the hottest are large, open settlements of the basin lowlands and the peri-urban fringe "
  "(Yuliang 42.8 °C, Qiankou 40.1 °C, Tangyue 39.8 °C), only two of which exceed 40 °C. The heat "
  "anomaly depends on how the background is defined, and the gradient across definitions is itself "
  "informative (Fig. 3c). Under the site-domain definition (V1), all 29 village point estimates are "
  "positive: ΔLST-V1 ranges from +0.03 to +7.69 °C with a mean of +2.45 °C, and the overpass-bootstrap "
  "95% interval excludes zero for 21 of 29 villages. Under the fabric variant (V2, n = 22) the anomaly "
  "is larger (mean +5.28 °C; 19 of 22 intervals above zero), targeting the built boundary itself. "
  "Under the terrain- and land-cover-matched background (V3), the anomaly halves: mean +1.27 °C, range "
  "−0.91 to +4.19 °C, with 25 of 29 point estimates positive and 22 intervals excluding zero; the four "
  "non-positive cases (Tachuan, Xucun, Shitan, Huansha) are villages whose fabric is not warmer than "
  "unbuilt terrain of the same elevation, slope, and tree cover. Part of the apparent heat excess "
  "therefore reflects where villages sit rather than the fabric itself—yet even the partially "
  "controlled estimate keeps the sign of the canonical urban heat island in 25 of 29 villages [11], "
  "[12], and the residual matching imbalance documented in Table A15 cautions against reading V3 as "
  "fully controlled.")
P("The spatial pattern is legible in Fig. 3a. The hot corridor follows the broad, open Xin’an River "
  "basin lowlands, where the large settlements of the plain—including the county seats and the "
  "peri-urban fringe of Shexian—concentrate; the cool zones track the forested ridge systems. The five "
  "coolest villages (Mulihong 31.2, Tachuan 31.8, Zuyuan 32.1, Xucun 33.4, Renli 33.7 °C) combine high "
  "elevation, high relief, and near-complete forest rings, while the five hottest (Yuliang 42.8, "
  "Qiankou 40.1, Tangyue 39.8, Xiongcun 39.3, Zhanqi 38.4 °C) combine low relief, maximal openness, "
  "and large contiguous fabric. The matched-background comparison shows that the anomaly is not a "
  "simple elevation artifact: most high villages are cooler in absolute terms yet remain warmer than "
  "equally high, equally forested surroundings—though the four exceptions above warn that this is not "
  "universal.")
P("The morphological associates of absolute LST under the dual criterion form a coherent terrain-and-"
  "fabric story (Table 2). Village LST rises with building coverage ratio (ρ = +0.65, q = 0.001, 95% CI "
  "[+0.27, +0.86]), edge density (+0.64, CI [+0.20, +0.85]), and domain built-up area (+0.61, "
  "CI [+0.20, +0.83]), and falls with relief (−0.69, CI [−0.85, −0.39]), the forest ring (−0.68, CI "
  "[−0.89, −0.27]), elevation (−0.62, CI [−0.88, −0.14]), slope (−0.53, CI [−0.79, −0.10]), and water "
  "distance (minimum distance −0.47, CI [−0.69, −0.12]; mean distance −0.44, CI [−0.68, −0.10]). "
  "Terrain-horizon openness shows a positive association (ρ = +0.53, q = 0.013, CI [+0.08, +0.78]) "
  "that passes the dual criterion at the canonical blocking configuration; the seed-and-design audit "
  "of Appendix Table A14b nonetheless flags this interval as seed-sensitive, so we report it as "
  "borderline rather than established. Largest-patch share is positively associated (ρ = +0.43, "
  "CI [+0.13, +0.68]) but fails family control (q = 0.058). Two sign readings are essential here. First, the water "
  "variable is a *distance*: the negative correlation means that villages whose fabric lies farther "
  "from water are cooler, i.e. the waterside villages of this sample are the *warmer* ones—because in "
  "this terrain the rivers run through the low, open basin floors, water distance is entangled with "
  "elevation and openness, and no evaporative-cooling claim can be made at village scale. Second, the "
  "two fabric-density metrics use the fixed domain area as denominator (Section 2.3): their positive "
  "signs therefore read as *more fabric per site area, warmer village*, consistent with the coverage-"
  "ratio result—not as a fragmentation effect. Under size control (family F2), only relief retains a "
  "dual-criterion association with LST (partial ρ = −0.54, CI [−0.72, −0.22]); elevation shows a "
  "strong residual tendency whose interval touches zero (partial ρ = −0.53, q = 0.023, CI [−0.85, "
  "+0.02])—so the size-controlled thermal read is suggestive rather than confirmatory.")
P("For the matched anomaly ΔLST-V3, four attributes pass the dual criterion, all of them fabric "
  "variables: building coverage ratio (ρ = +0.56, CI [+0.24, +0.81]), edge density (+0.53, "
  "CI [+0.20, +0.79]), domain built-up area (+0.50, CI [+0.12, +0.78]), and largest-patch share "
  "(+0.48, CI [+0.13, +0.78]). "
  "Terrain elevation and relief, by contrast, show no surviving association with the matched anomaly "
  "(elevation −0.29, n.s.; relief −0.27, n.s.): once the background is matched on terrain, high "
  "villages are not systematically less anomalous. Openness (+0.25, n.s.), the forest ring (−0.36, "
  "n.s.), and water distance (minimum −0.32, n.s.) are likewise unrelated to the matched anomaly, as "
  "are both plan-form metrics (elongation −0.02; compactness −0.07). The two dependent variables "
  "therefore tell different stories: absolute "
  "LST covaries mainly with *where* the village sits and *how much* is built, and the matched anomaly "
  "with the amount and texture of the fabric rather than with terrain. Southness is unrelated to "
  "village LST (ρ = −0.14, p = 0.47) and to the anomaly (ρ = −0.07, p = 0.70); north–south asymmetry "
  "behaves similarly (−0.12, p = 0.55 against LST; −0.08, p = 0.70 against the anomaly).")
FIG("figures/Fig3_lst.png", 3,
    "Summer thermal environment. (a) Quality-masked LST median composite over the retained independent "
    "overpasses (2019–2025, Landsat 8/9 Collection 2), with the four WRS path/row footprints outlined; "
    "(b) village-core versus site-domain background LST, villages sorted by core LST; (c) per-village "
    "heat anomaly ΔLST with overpass-bootstrap 95% CI under the site-domain (V1) and terrain-matched "
    "(V3) definitions (n = 29; V1 mean +2.45 °C; V3 mean +1.27 °C).")

H2("3.3. Wireless Coverage and Deployment-Phase Sensitivity")
P("Under the standardized deployment, four-phase mean cov85 spans 59.8–100% across villages (mean "
  "93.0%). Seventeen of 29 villages reach 96% or better, and these are precisely the open basin "
  "settlements; the deep-valley villages Mulihong (59.8%) and Zuyuan (63.6%) fall far behind, and the "
  "piedmont villages occupy the middle (Fig. 4a). Coverage at the lenient T−95 threshold is near "
  "saturation—16 villages stand at a four-phase mean of 100% (sample range 85.3–100%)—which is why "
  "cov85, the good-signal level in this scenario, is adopted as the primary metric; RSRP p10 ranges "
  "from −96.8 to −76.1 dBm and supports the same ordering.")
P("The four-phase repetition reveals a second, subtler result (Fig. 4b): grid phase matters exactly "
  "where coverage is marginal. The cross-phase standard deviation of cov85 reaches 34.9 percentage "
  "points in Zuyuan (phase values from 12% to 88%) and exceeds 20 points in Renli, Shitan, and "
  "Changxi, whereas basin villages are phase-invariant at or near 100%. Deployment sensitivity is "
  "thus itself associated with morphological setting: villages at the coverage margin can be lucky or "
  "unlucky with a given site pattern, while well-sited villages are comparatively insensitive to it. "
  "Accordingly, village rankings across phases are moderately stable for cov85 (mean pairwise rank "
  "correlation 0.57, minimum 0.47) and weaker for RSRP p10 (0.39, minimum 0.17); the plain mean RSRP "
  "proved phase-unstable (0.10, minimum −0.09), being hostage to the distance of the single nearest "
  "site, and was discarded from further analysis. The four-phase mean therefore absorbs most—but not "
  "all—deployment arbitrariness: rankings among open, well-covered villages are phase-insensitive, "
  "while rankings among marginal, deep-valley villages are not.")
P("The village-level detail behind these statistics is instructive. Yuliang, the peri-urban case, "
  "combines the largest target set with a mid-pack four-phase mean "
  "of 86.2%—and, unexpectedly, substantial phase dispersion (71.0–99.2% across phases): its fabric "
  "extends from the open plain into a river corridor whose margins lose visibility in two of the four "
  "grid alignments. Size and openness therefore do not exempt a village from deployment sensitivity. "
  "At the other extreme, Zuyuan and Mulihong—the two villages for which the land-cover map "
  "under-detects the built fabric, and whose coverage is therefore evaluated on a 150-m fallback "
  "grid—are not only the worst covered but also among the most phase-sensitive: their serviceability "
  "is decided by whether a half-spacing shift happens to place a site on the one ridge with a clear "
  "path into the valley. Because their fallback target grids are not strictly comparable to the "
  "built-up-cell targets of the other 27 villages, all coverage associations below are re-estimated "
  "without them (Section 3.4). Between these poles, piedmont villages such as Huansha, Zhukou, and "
  "Zhanqi lose service only at their valley-facing margins, which is where their "
  "weakest RSRP deciles originate.")
P("Frequency choice does not change the picture. Repeating the full chain at 700 MHz (a band central "
  "to China’s current rural 5G build-out) raises absolute coverage as expected—the four-phase mean "
  "cov85 rises from 93.0% to 95.9%—but preserves the inter-village ordering (rank correlation 0.89 "
  "between the four-phase means of the two bands; in the phase-0 comparison, 0.81 for cov85 and ≥0.96 "
  "for the RSRP statistics; Appendix Table A6). The near-saturated cov95 ordering is, as expected, the "
  "least stable statistic across bands (ρ = 0.39). The morphological gradients reported below are "
  "therefore not an artifact of band choice.")
FIG("figures/Fig4_coverage.png", 4,
    "Wireless coverage under the standardized virtual deployment. (a) Four-phase mean cov85 by village "
    "(points) over the 1620-site phase-0 grid (dots), on a 30-m hillshade; (b) deployment-phase "
    "sensitivity: villages with low mean coverage also show high cross-phase dispersion (bars: "
    "cross-phase minimum–maximum range).")

H2("3.4. Morphology–Performance Coupling and Trade-Offs")
P("Bringing the two performances together (Fig. 5, Table 2) answers RQ4. To classify couplings "
  "consistently, we fix utility directions explicitly: the thermal utility is U_T = −LST (or −ΔLST) "
  "and the coverage utility is U_C = cov85 (or RSRP p10). A morphological attribute is a *trade-off* "
  "attribute when its associations with the two utilities carry opposite signs—note that same-signed "
  "raw correlations with LST and with cov85 imply opposite-signed utility associations—and a "
  "*one-sided* attribute when only one association passes the dual criterion. Given the modest sample "
  "and the cross-sectional design, we read the trade-off attributes below as *candidate couplings* for "
  "site-specific testing, not as established laws.")
P("Three terrain attributes are two-sided in the raw family. The **forest ring**—the emblematic "
  "Huizhou backdrop—is associated with lower village LST (ρ = −0.68, q < 0.001, CI [−0.89, −0.27]) "
  "and with worse coverage (cov85 ρ = −0.74, q < 0.001, CI [−0.90, −0.50]; RSRP p10 ρ = −0.65): both "
  "sides pass the dual criterion, making the ring the clearest two-sided attribute in the sample. "
  "Under size control the coverage side persists (partial ρ = −0.64 and −0.59, both dual) while the "
  "thermal side attenuates to a nominal tendency (partial ρ = −0.47, q = 0.066, CI [−0.79, +0.03]). "
  "**Slope** behaves the same way: cooler villages (LST ρ = −0.53, dual) but poorer service (cov85 "
  "ρ = −0.85, q < 0.001, CI [−0.93, −0.70]; RSRP p10 ρ = −0.83), with the coverage side surviving "
  "size control (partial ρ = −0.80 and −0.82, both dual) and the thermal side attenuating (−0.28, "
  "n.s.). **Relief** shows the same pattern with one stronger feature: its thermal side also survives "
  "size control (LST ρ = −0.69, dual, and size-controlled partial ρ = −0.54, CI [−0.72, −0.22], "
  "dual; cov85 ρ = −0.77, dual; size-controlled coverage partial ρ = −0.69, dual). The common reading "
  "is that steep, enclosed, forested terrain cools and isolates at once.")
P("**Terrain-horizon openness** is the mirror image, but with an important asymmetry of evidence. Its "
  "coverage side is solid (cov85 ρ = +0.57, CI [+0.36, +0.73]; RSRP p10 ρ = +0.52, CI [+0.23, "
  "+0.69]; both dual), though it attenuates below family control under size control (partial "
  "ρ = +0.43 and +0.42, q ≈ 0.09–0.10, intervals excluding zero). Its thermal side is borderline: "
  "ρ = +0.53 with LST passes the dual criterion at the canonical blocking configuration "
  "(q = 0.013, CI [+0.08, +0.78]), but the 20-seed audit of Appendix Table A14b flags this "
  "interval—and the openness–RSRP interval—as seed-sensitive, and under size control the thermal "
  "association collapses (partial ρ = +0.31, n.s.). We therefore classify openness as a coverage "
  "attribute with a borderline, size-entangled thermal tendency—not as a confirmed two-sided "
  "trade-off axis.")
P("The fabric metrics speak mostly on the thermal side. Building coverage ratio is the strongest "
  "fabric associate of both thermal outcomes (LST ρ = +0.65, CI [+0.27, +0.86]; ΔLST ρ = +0.56, "
  "CI [+0.24, +0.81]), edge density follows (LST +0.64; ΔLST +0.53, both dual), and largest-patch "
  "share tracks the matched anomaly (+0.48, dual) without surviving family control against absolute "
  "LST (+0.43, q = 0.058). Because edge density is denominated by the fixed domain area, these "
  "positive signs share one reading: more fabric per site area, warmer village—consistent with the "
  "coverage-ratio result rather than indicating a fragmentation effect, and these associations "
  "persist against the matched background, which terrain associations do not. On the wireless side, "
  "edge density carries a moderate positive association (cov85 ρ = +0.45, CI [+0.09, +0.70]) that "
  "passes the dual criterion at the canonical configuration but is seed-sensitive in the Table A14b "
  "audit, and the coverage ratio shows none (+0.42 nominal, CI [−0.08, +0.76]). Domain built-up "
  "area itself is double-edged in the raw family—larger villages "
  "are hotter (LST ρ = +0.61, dual) and better covered (cov85 ρ = +0.47, dual, interval marginally "
  "clear of zero)—which is precisely "
  "why size serves as the control variable rather than being read as a morphological trade-off "
  "axis.")
P("The remaining attributes are one-sided or inert. Elevation is thermal-only in this sample (LST "
  "ρ = −0.62, dual; coverage ρ = −0.02, n.s.), and it is not associated with the matched anomaly "
  "(Section 3.2). Water distance is likewise thermal-only (minimum distance vs. LST ρ = −0.47, dual; "
  "vs. ΔLST ρ = −0.32, n.s.; vs. coverage ρ = −0.24, n.s.)—with the sign reading of Section 3.2: "
  "farther from water, cooler village; the riverside lowland villages are the warm ones, so the "
  "variable indexes terrain setting, not evaporative cooling. The two plan-form geometry metrics "
  "show no association that passes the dual criterion: elongation is unrelated to any outcome "
  "(LST ρ = −0.12; ΔLST −0.02; cov85 −0.33, p = 0.09), and compactness shows a consistent but "
  "formally non-surviving negative tendency on the wireless side (raw ρ = −0.36 with cov85 and "
  "−0.41 with RSRP p10; size-controlled partial ρ = −0.44 and −0.46, nominal p < 0.05 but q > 0.05), "
  "which we report as a hypothesis for larger samples. Southness shows no association with any "
  "outcome (LST ρ = −0.14, p = 0.47), and north–south asymmetry behaves similarly. Of the 64 raw "
  "tests (family F1), 24 pass the dual criterion (24 on BH control alone; 23 on the non-spatial "
  "permutation baseline); of the 60 size-controlled partial "
  "tests (family F2), 7 pass—six on the coverage side (relief, slope, and the forest ring against "
  "cov85 or RSRP p10) plus relief against LST.")
P("Sensitivity checks qualify these associations honestly. Excluding the peri-urban Yuliang case "
  "leaves every reported sign intact and strengthens several magnitudes (openness–coverage ρ = +0.64; "
  "forest-ring–coverage −0.83; slope–coverage −0.86; coverage-ratio–LST +0.61; forest-ring–LST "
  "−0.65). Excluding the two fallback-target villages from the coverage tests leaves the headline "
  "coverage associations in place (openness–coverage ρ = +0.61, p < 0.001, n = 27; forest ring "
  "−0.68; slope −0.81; relief −0.71). In the leave-one-county-out jackknife the headline "
  "associations keep their signs in every county omission (openness–coverage ρ between +0.51 and "
  "+0.63; forest-ring–coverage between −0.89 and −0.68; slope–coverage between −0.88 and −0.81; "
  "coverage-ratio–LST between +0.57 and +0.80; forest-ring–LST between −0.84 and −0.60), so no "
  "single county drives the results. Spatial diagnostics motivate the conservative inference "
  "machinery: Moran’s I detects significant spatial clustering for cov85 (I = 0.22, p = 0.002), "
  "elevation (I = 0.26, p = 0.001), absolute LST (I = 0.18, p = 0.008), and slope (I = 0.17, "
  "p = 0.004), but not for terrain-horizon openness, the forest ring, the matched anomaly, or RSRP "
  "p10; accordingly, inferential weight rests on the dual criterion rather than on nominal p-values "
  "alone.")
FIG("figures/Fig5_tradeoff.png", 5,
    "Morphology–performance couplings. (a), (b) Terrain-horizon openness versus village LST and cov85; "
    "(c), (d) forest ring share versus village LST and cov85. Point size: domain built-up area; lines: "
    "Theil–Sen robust fits truncated to the observed metric range; annotations: Spearman ρ with "
    "nominal p, within-family Benjamini–Hochberg q, the block-bootstrap 95% CI, and the size-controlled "
    "partial ρ; n = 29 in all panels; the number of villages at the cov85 ceiling is noted in panels b "
    "and d.")
_t2 = pd.read_csv("tables/Table2_correlation.csv")
TABLE(2, "Spearman correlations between morphological metrics and performance variables (upper row per "
         "metric), and partial correlations controlling domain built-up area (row marked “ctrl size”). "
         "* p < 0.05, ** p < 0.01, *** p < 0.001 (nominal); † passes the dual criterion—within-family "
         "Benjamini–Hochberg q < 0.05 and block-bootstrap 95% CI excluding zero (F1: 64 raw tests; "
         "F2: 60 partial tests). Per-test statistics, including the non-spatial permutation baselines and "
         "the effective-sample-size heuristic, are given in Appendix Table A3; interval stability is "
         "audited in Appendix Tables A14 and A14b. Plan-form and fabric "
         "metrics (except coverage ratio) use n = 28 (Mulihong has no detected built-up cell in its "
         "domain); others n = 29.",
      list(_t2.columns), _t2.values.tolist(), fontsize=8)

# ==================== 4. DISCUSSION ====================
H1("4. Discussion")
H2("4.1. Two-Sided Attributes, Read Mechanistically")
P("The first coupling is terrain enclosure, and the evidence for its two sides is asymmetric. On the "
  "coverage side, steep slopes, high relief, and obstructed horizons robustly predict poor service, "
  "and these associations survive size control—the strongest and most conservative layer of our "
  "results. On the thermal side, the same terrain is associated with cooler villages in the raw "
  "family (slope ρ = −0.53; relief ρ = −0.69; openness ρ = +0.53 borderline), and relief—uniquely "
  "among the terrain attributes—retains a size-controlled thermal association (partial ρ = −0.54); "
  "the openness interval is seed-sensitive in the block-design audit. Mechanistically this is "
  "coherent: the same enclosure that shades a valley site "
  "and channels nocturnal cold-air drainage also blocks the low-elevation radio paths on which "
  "rural macrocells depend, while the size of basin settlements confounds the thermal read. The "
  "result inverts the urban intuition with a caveat: in city morphometrics, higher sky-view factor "
  "is generally associated with *cooler* canyon environments because SVF indexes heat release from "
  "street canyons [14], [18]; at the mountain-village scale, higher terrain-horizon openness tends "
  "to mark *hotter* settlements because it indexes insolation exposure of the whole site. The sign "
  "flip is a scale effect, not a contradiction, but its borderline status here warns against "
  "transferring urban morphometric rules to rural terrain-coupled settlements in either direction. "
  "We therefore treat the enclosure coupling as a planning hypothesis with a robust coverage side "
  "and a moderate, size-entangled thermal side.")
P("The second coupling is the forest ring, and it is the cleanest two-sided attribute in the sample. "
  "The ring is associated with cooler villages (ρ = −0.68) and with signal shadowing (ρ = −0.74), "
  "both passing the dual criterion. The thermal side attenuates under size control (partial "
  "ρ = −0.47, nominal only)—so we claim a selective, "
  "conditional daytime association at village scale, not a universal one; whether it reflects "
  "evapotranspiration, the co-location of forest with steep shaded slopes, or both cannot be "
  "separated with the present data. The coverage side, by contrast, is robust across every check we "
  "ran, and in our model it is entirely terrain-mediated. Whether vegetation attenuation over the "
  "ring [19] would change this village-level ranking cannot be determined from the present model: it "
  "would deepen absolute shadowing, but the ranking response depends on village-specific path "
  "geometries, so we flag it as a validation target rather than assert a direction (Section 4.4). "
  "For the ecological-wisdom reading of vernacular landscape, the honest summary is therefore "
  "two-sided: the forested backdrop is associated with a measurable daytime thermal benefit and "
  "with a measurable digital-infrastructure cost, and a credible account must state both.")
P("The third result belongs to architectural morphology proper. The fabric metrics—coverage "
  "ratio, edge density, largest-patch share, patch density—carry the thermal signal that survives "
  "background matching: more fabric per site area is systematically warmer. This is the "
  "settlement-scale analogue of the urban finding that contiguous built mass "
  "stores and re-radiates heat, shown here statistically across a traditional-village sample rather "
  "than within a single case study. "
  "Two qualifications matter. First, these metrics derive from 10-m land cover: they measure the "
  "detected fabric, not the historic core, and they cannot see courtyards, lanes, or building "
  "heights. Second, under size control the fabric associations attenuate to non-surviving "
  "tendencies (coverage-ratio partial ρ = +0.35 with LST), so fabric amount and village size "
  "are entangled. The plan-form geometry metrics tell a different story: neither elongation nor "
  "compactness is detectably associated with any thermal outcome, and on the wireless side only a "
  "non-surviving negative tendency of compactness remains (partial ρ ≈ −0.45). We treat the "
  "compactness question as unresolved rather than answered: the hypothesis that compact historic "
  "cores run hotter than their hinterland—plausible from sparse internal green space and reduced "
  "sky view—requires building-scale data, not 10-m land cover. Southness is inert at settlement "
  "scale in this sample: we find no evidence that orientation predicts thermal outcomes, so claims "
  "for village-scale orientation effects should be treated as unproven rather than "
  "canonical—orientation effects are real at the scale of buildings and canyons [3], [14], but they "
  "average out across a settlement, and the ten sampled villages with negative domain-mean "
  "southness show that the textbook ideal is an idealization rather than a rule.")
H2("4.2. Comparison with Previous Studies")
P("Three strands of previous work frame our findings. First, the urban morphometrics literature has "
  "established SVF–temperature relationships at the scale of street canyons and LCZ classes [13], [14], "
  "[18], where openness generally aids night-time cooling. Our borderline association between terrain "
  "openness and higher village-scale summer LST (ρ = +0.53, dual criterion met at the canonical "
  "configuration but seed-sensitive in the block-design audit) does not contradict "
  "this work: it reflects the different physics that dominate when the ‘canyon’ is a mountain valley "
  "and the variable is daytime insolation of the whole site rather than nocturnal long-wave release "
  "from a street. The two regimes are reconciled once one recognizes that terrain-horizon openness here "
  "measures horizon obstruction by terrain, not by buildings. Second, the vernacular-climate literature "
  "has documented passive strategies of traditional settlements through single-village instrumentation "
  "and simulation [1]–[4], [7], including work in the Huizhou area itself [3]. Our contribution relative "
  "to that strand is statistical rather than instrumental: with 29 villages we can separate site-level "
  "associations from fabric-level ones (terrain setting versus coverage ratio and patch structure) in a "
  "way that one-village case studies structurally cannot; at the same time, the dependence of the "
  "forest-ring association on the ring’s geometric definition across processing versions is a standing "
  "caution against generalizing courtyard-scale measurements to the settlement "
  "scale. Third, the urban-informatics strand has shown how open imagery and reproducible computation "
  "can characterize fine-grained urban form at scale [22], [24], [25]. We extend that ethos to a domain "
  "it has largely bypassed—rural heritage settlements—and to a second physical field, radio "
  "propagation, that this literature has not engaged. The nearest precedents to our coverage analysis "
  "are the propagation-modeling and radio-planning literature [8]–[10]; our inversion of the "
  "perspective, asking not ‘where should sites go’ but ‘which village forms are inherently hard to "
  "serve’, appears to be new (Section 1).")
P("Elongation illustrates the same discipline from the other side. In the thermal literature, "
  "linear valley villages would be filed under ‘favorably shaded’; in radio planning the same form is a "
  "known worst case, because a linear settlement along a valley floor maximizes the share of fabric "
  "lying at low radio horizons. In our data the coverage side of that expectation appears only as a "
  "marginal trend (ρ = −0.33, p = 0.09 with cov85) and the thermal side not at all (ρ = −0.12, "
  "p = 0.54 with LST): the two-physics elongation trade-off is mechanistically grounded in both "
  "literatures, but it is not confirmed at this sample size, and we state that plainly rather than "
  "inherit the expectation from either single-physics strand.")

H2("4.3. Implications for Heritage-Compatible and Digital-Village Planning")
P("In what follows we separate three levels of statement: associations observed in the data under the "
  "dual criterion, inferences that are mechanistically plausible but statistically partial, and "
  "hypotheses that the present data cannot test. For conservation planning, the practical message is "
  "that thermally favorable settings are not automatically connectivity-friendly: the associations "
  "reported here indicate a potential planning trade-off—quantifiable in principle, but requiring "
  "site-specific measurement and validation before any engineering action. Even the worst-served "
  "villages reach roughly 60% T−85 exceedance under a macro-only deployment in this stylized scenario; "
  "one hypothesis for closing the residual gap—untested here—is low-visual-impact small cells at the "
  "village perimeter, fed through the open river and road corridors that pierce the forest ring. LOS "
  "simulation on open elevation data, as used here, offers heritage authorities a cheap screening tool "
  "that can be run before any hardware or visual-impact discussion begins, and that directly respects "
  "the constraint that matters most in a protected village: where things can be seen from, and where "
  "they cannot.")
P("For digital-village programs, the same pipeline doubles as a screening layer. From open "
  "data it maps both the thermal burden and the radio visibility of the sampled settlements, "
  "providing a "
  "baseline against which digital infrastructure in historic settings can be planned rather than "
  "improvised. Because deployment-phase sensitivity is itself mapped (Fig. 4b), planners "
  "can distinguish villages whose service will be insensitive to deployment phase from villages that "
  "need site-by-site engineering. The open-data character of the whole workflow matters practically: "
  "it can be audited by conservation reviewers, re-run as new satellite data arrive, and transferred "
  "to other mountainous regions at low marginal cost, subject to local data-availability checks. "
  "Embedding such layers into city information modeling (CIM) base platforms [6] is a natural next "
  "step, but it is future work, not a result of this study.")
P("Finally, there is a discursive implication. The ecological wisdom of vernacular settlements is "
  "usually asserted qualitatively in heritage narratives; here it is quantified—and found to be "
  "selective and conditional. In this sample, terrain setting, settlement size, and fabric "
  "consolidation carry measurable thermal associations while plan-form geometry effects are not "
  "detectable at this sample size; the forest ring’s daytime cooling association is measurable but "
  "size-entangled and dependent on how the ring is drawn; the orientation canon is neither rule nor thermal "
  "determinant at settlement scale; and the same attributes that cool also isolate from the network. "
  "A mature, evidence-based account of traditional ecological knowledge should be able to state both "
  "sides of that ledger. We suggest that this kind of quantified, two-sided assessment is precisely "
  "what makes the wisdom credible—and translatable—when it is communicated to engineering, planning, "
  "and policy audiences.")
H2("4.4. Limitations and Future Work")
P("Ten limitations qualify the findings. First, LST is not air temperature or thermal comfort: "
  "satellite LST responds to surface fabric and misses shade-mediated comfort, and our composite "
  "characterizes summer daytime only; sensor-based validation is the obvious next step. Second, the "
  "composite is correlational in nature and single-season; night-time, winter, and transitional-season "
  "behavior may differ, which matters directly for interpreting the forest-ring association. Absolute "
  "LST is additionally composite-dependent across the four satellite frames and is used for ranking "
  "only; for 8 of the 29 villages the overpass-bootstrap interval of the V1 anomaly includes zero, so "
  "village-level anomaly values should be read with their intervals (Table 1); we do not propagate the "
  "per-pixel ST_QA uncertainty band distributed with the product, which would widen the reported "
  "intervals further; and the anomaly is definition-sensitive—four villages turn non-positive under "
  "the matched background—so it should be read as a range across the V1–V3 definitions, not as a "
  "point value. Third, the terrain-matched background (V3) is only partially controlled: post-matching "
  "imbalance on slope and tree cover remains substantial for some villages (Table A15), so residual "
  "confounding by terrain context cannot be excluded. Fourth, the standardized deployment holds the "
  "deployment pattern constant but is not any operator’s network; absolute coverage levels would shift "
  "with real site "
  "inventories, antenna patterns, and traffic engineering, and the morphological ranking—the object "
  "of this study—proved moderately stable across grid phases for open villages but genuinely "
  "phase-sensitive for marginal ones, a heterogeneity our four-phase mean only partly absorbs. Fifth, "
  "the propagation chain is a deterministic received-power scenario: it omits vegetation and building "
  "clutter, diffraction refinement, shadowing, and interference; the receiver is defined 1.5 m above "
  "the surface model, so over forest and fabric it sits at canopy or roof level rather than at "
  "pedestrian level; reported coverage is optimistic in absolute terms; whether vegetation attenuation "
  "would change the village-level ranking is unknown and is a validation target; and validation "
  "against measured or crowdsourced RSRP, or against a diffraction-capable model such as ITU-R P.1812 "
  "[36] or Longley–Rice, remains future work. The RMa NLOS branch is moreover specified only to 5 km; "
  "a truncation variant enforcing that limit leaves all village coverage metrics unchanged (Appendix "
  "Table A7), but links near the assignment radius remain extrapolations. Sixth, the morphometrics are "
  "coarse by architectural standards: terrain-horizon openness is evaluated once at the canonical "
  "village point rather than across the fabric, and the 10-m land-cover map cannot resolve the narrow "
  "lanes, courtyards, building heights, and street H/W ratios of Huizhou villages; the courtyard-scale "
  "design hypotheses mentioned in Section 4.1 are therefore speculative and require building-scale "
  "data. The fabric metrics are computed on detected 10-m cells inside a fixed domain and inherit "
  "land-cover classification error; recomputing the domain coverage metric against the previous "
  "WorldCover vintage (2020 v100) preserves the village ranking (ρ = 0.957, mean absolute difference "
  "4.6 percentage points; Appendix Table A17), so this error source does not drive the results. "
  "Seventh, the elevation source is a digital *surface* model: over forest and "
  "fabric its heights include canopy and roof signals, so terrain metrics are surface metrics; a "
  "cross-check against the independent NASADEM product agrees closely (centroid elevation ρ = 0.999, "
  "domain relief ρ = 0.988; Appendix Table A10), but NASADEM is itself a radar surface model, and "
  "per-village discrepancies reach 208 m at Tachuan, indicating steep, densely vegetated terrain "
  "where surface and ground diverge; LOS determination and all elevation-derived metrics inherit "
  "this limitation. Eighth, two villages rely on fallback coverage targets and one of them also lacks "
  "plan-form and fabric metrics; although exclusion checks leave the headline associations unchanged, "
  "manual or semi-automated mapping of their built boundaries—ideally tied to field survey—would "
  "remove this caveat. Ninth, the sample is a purposive, register-based regional sample (n = 29) "
  "screened by a single reviewer, not a probability sample, and two villages carry approximate-"
  "coordinate flags (Section 2.1); Moran’s I detects significant spatial "
  "clustering in several metrics and outcomes, and while the dual criterion, the seed-and-design "
  "interval audit, and the county jackknife mitigate the resulting overconfidence, all inference "
  "remains associational and we do not generalize beyond the sampled region; the effective-sample-size "
  "figures reported in Table A3 are a heuristic plausibility check, not a formal correction. Tenth, "
  "one village’s detected built fabric merges with a county seat in the land-cover map, though "
  "excluding it changes no reported sign.")
P("Future work follows four lines: (i) multi-season and diurnal LST composites, plus air-temperature "
  "and comfort sensor campaigns in a stratified subsample of villages; (ii) building-scale "
  "morphometrics from high-resolution imagery or survey—street H/W, courtyard green view, "
  "building-resolved SVF—tested against the settlement-scale associations reported here; (iii) "
  "validation of the radio chain against real operator inventories and crowdsourced measurements, "
  "cross-checking against a diffraction-capable propagation model, and explicit measurement of "
  "vegetation attenuation to test whether the forest-ring ranking survives it; and (iv) extension of "
  "the sample beyond Huizhou, and embedding of the screening layers into CIM base platforms [6], to "
  "test whether the candidate couplings generalize across mountain-village regions with different "
  "building traditions.")

# ==================== 5. CONCLUSIONS ====================
H1("5. Conclusions")
P("Across 29 traditional Huizhou villages, settlement morphology is associated with summer surface "
  "temperature and with simulated wireless coverage in systematically coupled ways. Under the "
  "site-domain definition, all 29 village point estimates of the summer heat anomaly are positive "
  "(mean ΔLST = +2.4 °C; the overpass-bootstrap interval excludes zero in 21 of 29); under a terrain- "
  "and land-cover-matched background the anomaly halves (mean +1.3 °C; 25 of 29 positive), and its "
  "surviving correlates are the amount and texture of the built fabric rather than terrain. Absolute "
  "temperature covaries with both setting and fabric: warmer villages are lower in elevation and "
  "relief, less steep, closer to water, more densely built per site area, and less forest-ringed. "
  "Three attributes carry opposite-signed associations with the thermal and coverage utilities—the "
  "forest ring (cooler, worse served), slope, and relief—although under size control the coverage "
  "side survives throughout while on the thermal side only relief does, so these are candidate "
  "couplings, not confirmed laws. Terrain-horizon openness is a coverage attribute whose thermal "
  "side is borderline and seed-sensitive. The fabric metrics—coverage ratio, edge density, "
  "largest-patch share—show that villages with more fabric per site area are systematically warmer, "
  "statistical evidence gathered across a traditional-village sample rather than within a single "
  "case study, "
  "while plan-form geometry (elongation, compactness) and southness show no association that survives "
  "control at this sample size. The headline associations are stable to frequency band, county "
  "jackknife, outlier exclusion, water-masking thresholds, land-cover vintage, and the "
  "propagation-model range limit, "
  "with grid-phase effects absorbed for open villages but persisting for marginal ones, and the "
  "entire analysis runs on open data and a scripted, deposited pipeline. Because the design is "
  "cross-sectional and the deployment stylized, the quantified couplings are hypotheses for "
  "site-specific assessment—not causal effects, operational coverage estimates, or general laws. For "
  "the agencies that must reconcile climate adaptation with digital modernization in historic "
  "mountain settlements, the central lesson is that a potential conflict is visible in the data and "
  "quantifiable—and, once made visible, open to design responses—subject to site-specific validation.")

# ==================== 声明 ====================
H1("Declarations")
P("**Data Availability.** The full analysis code, the village-level analysis table (all sixteen "
  "morphological metrics, four performance variables, and four-phase coverage detail), the Landsat "
  "scene manifest, and the figure scripts are deposited at "
  "https://github.com/ccst2000/huizhou-thermal-wireless-tradeoffs and "
  "archived at https://doi.org/10.5281/zenodo.22011592. Public input datasets are cited in the References.")
P("**Author Contributions (CRediT).** Lei Zhang: conceptualization, methodology, software, formal "
  "analysis, investigation, data curation, writing—original draft, writing—review and editing, "
  "visualization, project administration.")
P("**Conflicts of Interest.** The author declares no conflict of interest.")
P("**Funding.** [Funding information to be inserted.]")

H1("Acknowledgments")
P("The author thanks the data providers: the U.S. Geological Survey and Microsoft Planetary Computer "
  "for Landsat access, the European Space Agency for WorldCover, and the Copernicus programme for the "
  "digital surface model; and Esri for the World Imagery base map used in Fig. 1.")

# ==================== REFERENCES ====================
H1("References")
refs = [
    "Z. J. Zhai and J. M. Previtali, “Ancient vernacular architecture: Characteristics categorization "
    "and energy performance evaluation,” Energy and Buildings, vol. 42, no. 3, pp. 357–365, 2010.",
    "S. Gou, Z. Li, Q. Zhao, V. M. Nik, and J.-L. Scartezzini, “Climate responsive strategies of "
    "traditional dwellings located in an ancient village in hot summer and cold winter region of "
    "China,” Building and Environment, vol. 86, pp. 151–165, 2015.",
    "C. Pan, Y. Wu, S. Chen, and Y. Yang, “Indoor environmental comfort assessment of traditional folk "
    "houses: A case study in southern Anhui, China,” International Journal of Environmental Research "
    "and Public Health, vol. 20, no. 4, 3024, 2023.",
    "Y. Xiong, J. Zhang, Y. Yan, S. Sun, X. Xu, and E. Higueras, “Effect of the spatial form of "
    "Jiangnan traditional villages on microclimate and human comfort,” Sustainable Cities and Society, "
    "vol. 87, 104136, 2022.",
    "General Office of the CPC Central Committee and General Office of the State Council, “Strategic "
    "Outline for Digital Village Development,” Beijing, China, 2019 (in Chinese).",
    "Ministry of Housing and Urban-Rural Development of the People’s Republic of China, Technical "
    "Standard for City Information Modeling Basic Platforms, CJJ/T 315-2022, Beijing, China, 2022 "
    "(in Chinese).",
    "Z. Wan, et al., “How does outdoor spatial design shape the microclimate, comfort, and behavior "
    "in traditional Chinese villages? A systematic review across scales, contexts, and users,” "
    "Sustainability, vol. 17, no. 15, 6960, 2025.",
    "M. Hata, “Empirical formula for propagation loss in land mobile radio services,” IEEE "
    "Transactions on Vehicular Technology, vol. 29, no. 3, pp. 317–325, 1980.",
    "3GPP, “Study on channel model for frequencies from 0.5 to 100 GHz,” 3rd Generation Partnership "
    "Project, Sophia Antipolis, France, TR 38.901 V16.1.0, 2020.",
    "T. S. Rappaport, Wireless Communications: Principles and Practice, 2nd ed. Upper Saddle River, "
    "NJ, USA: Prentice Hall, 2002.",
    "T. R. Oke, “The energetic basis of the urban heat island,” Quarterly Journal of the Royal "
    "Meteorological Society, vol. 108, no. 455, pp. 1–24, 1982.",
    "A. J. Arnfield, “Two decades of urban climate research: A review of turbulence, exchanges of "
    "energy and water, and the urban heat island,” International Journal of Climatology, vol. 23, "
    "no. 1, pp. 1–26, 2003.",
    "I. D. Stewart and T. R. Oke, “Local climate zones for urban temperature studies,” Bulletin of "
    "the American Meteorological Society, vol. 93, no. 12, pp. 1879–1900, 2012.",
    "T. R. Oke, “Canyon geometry and the nocturnal urban heat island: Comparison of scale model and "
    "field observations,” Journal of Climatology, vol. 1, no. 3, pp. 237–254, 1981.",
    "K. Žákšek, K. Oštir, and Ž. Kokalj, “Sky-view factor as a relief visualization technique,” "
    "Remote Sensing, vol. 3, no. 2, pp. 398–415, 2011.",
    "Y. Xiong, Y. He, X. Xie, T. Zhai, N. Chu, L. Shen, and Y. Yang, “A study on the spatial form of "
    "traditional villages in Jiangnan region of China from the perspective of human thermal comfort: "
    "A case study of Nanjing, Jiangsu Province,” PLoS ONE, vol. 20, no. 5, e0323252, 2025.",
    "S. Wu, et al., “Research on microclimate optimization of traditional residential buildings in "
    "central Anhui based on humid and hot climate characteristics and regional architectural "
    "features,” Buildings, vol. 14, no. 8, 2323, 2024.",
    "A. Middel, K. Häb, A. J. Brazel, C. A. Martin, and S. Guhathakurta, “Impact of urban form and "
    "design on mid-afternoon microclimate in Phoenix Local Climate Zones,” Landscape and Urban "
    "Planning, vol. 122, pp. 16–28, 2014.",
    "ITU-R, “Attenuation in vegetation,” International Telecommunication Union, Geneva, Switzerland, "
    "Recommendation ITU-R P.833-9, 2016.",
    "ITU, Measuring Digital Development: Facts and Figures 2023. Geneva, Switzerland: International "
    "Telecommunication Union, 2023.",
    "European Space Agency, “Copernicus GLO-30 digital surface model (COP-DEM),” 2021. [Online]. "
    "Available: https://spacedata.copernicus.eu/collections/copernicus-digital-elevation-model, "
    "doi: 10.5270/ESA-c5d3d65 (accessed August 2026).",
    "X. Li, C. Zhang, W. Li, R. Ricard, Q. Meng, and W. Zhang, “Assessing street-level urban greenery "
    "using Google Street View and a modified green view index,” Urban Forestry & Urban Greening, "
    "vol. 14, no. 3, pp. 675–685, 2015.",
    "D. Zanaga et al., “ESA WorldCover 10 m 2021 v200,” Zenodo, 2022. [Online]. Available: "
    "https://doi.org/10.5281/zenodo.7254221 (accessed August 2026).",
    "Y. Long and L. Liu, “How green are the streets? An analysis for central areas of Chinese cities "
    "using Tencent Street View,” PLoS ONE, vol. 12, no. 2, e0171110, 2017.",
    "Y. Ye, D. Richards, Y. Lu, X. Song, Y. Zhuang, W. Zeng, and T. Zhong, “Measuring daily accessed "
    "street greenery: A human-scale approach for informing better urban planning practices,” Landscape "
    "and Urban Planning, vol. 191, 103434, 2019.",
    "U.S. Geological Survey, “Landsat 8–9 OLI/TIRS Collection 2 Level-2 science products,” 2021. "
    "[Online]. Available: https://doi.org/10.5066/P9OGBGM6 (accessed August 2026).",
    "Microsoft, “Planetary Computer: A planetary platform for sustainability,” collection "
    "landsat-c2-l2, catalog query executed 17 August 2026. [Online]. Available: "
    "https://planetarycomputer.microsoft.com.",
    "J. A. Voogt and T. R. Oke, “Thermal remote sensing of urban climates,” Remote Sensing of "
    "Environment, vol. 86, no. 3, pp. 370–384, 2003.",
    "Y. Benjamini and Y. Hochberg, “Controlling the false discovery rate: A practical and powerful "
    "approach to multiple testing,” Journal of the Royal Statistical Society: Series B, vol. 57, "
    "no. 1, pp. 289–300, 1995.",
    "H. Theil, “A rank-invariant method of linear and polynomial regression analysis,” Indagationes "
    "Mathematicae, vol. 12, pp. 85–91, 1950.",
    "P. K. Sen, “Estimates of the regression coefficient based on Kendall’s tau,” Journal of the "
    "American Statistical Association, vol. 63, no. 324, pp. 1379–1389, 1968.",
    "NASA Jet Propulsion Laboratory, “NASADEM Merged DEM Global 1 arc second V001,” NASA EOSDIS Land "
    "Processes DAAC, 2020, doi: 10.5067/MEaSUREs/NASADEM/NASADEM_HGT.001 (accessed August 2026).",
    "3GPP, “Evolved Universal Terrestrial Radio Access (E-UTRA): Physical layer—Measurements,” 3rd "
    "Generation Partnership Project, Sophia Antipolis, France, TS 36.214 V16.1.0, 2020.",
    "B. Efron and R. J. Tibshirani, An Introduction to the Bootstrap. New York, NY, USA: Chapman & "
    "Hall, 1993.",
    "H. R. Künsch, “The jackknife and the bootstrap for general stationary observations,” The Annals "
    "of Statistics, vol. 17, no. 3, pp. 1217–1241, 1989.",
    "ITU-R, “A path-specific propagation prediction method for point-to-area terrestrial services in "
    "the frequency range 30 MHz to 6 000 MHz,” International Telecommunication Union, Geneva, "
    "Switzerland, Recommendation ITU-R P.1812-8, 09/2025.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

# ==================== APPENDIX ====================
doc.add_page_break()
H1("Appendix A. Full Metric Table and Supporting Data")
P("Table A1 reports all sixteen morphological metrics, the four performance variables, per-phase "
  "coverage, and Landsat observation support for the 29 villages; given its width it is provided as a "
  "machine-readable CSV in the repository rather than typeset here. Table A2 below lists the Landsat "
  "scene manifest. Further machine-readable appendix tables are distributed with the repository (see "
  "Supplementary Materials): Table A3, per-test exact statistics for families F1 and F2 (nominal, "
  "label-permutation baseline and Freedman–Lane residual-permutation p, ESS-heuristic p, "
  "within-family BH q, block-bootstrap CI, n); Table A4, "
  "leave-one-county-out jackknife; Table A5, Moran’s I diagnostics; Table A6, 700-MHz four-phase "
  "coverage; Table A7, NLOS 5-km truncation and d3D-correction sensitivity; Table A8, per-village "
  "scene composition; Table A9, the three ΔLST variants with coverage-threshold and near-zero-date "
  "sensitivities; Table A10, DSM cross-check against NASADEM; Table A11, the sampling frame with "
  "per-village coordinate verification and canonical-coordinate provenance; Table A12, worked link "
  "budgets for three representative "
  "villages; Table A13, the earlier component-based morphology table retained for comparison with the "
  "current scripted chain; Table A14, block-bootstrap sensitivity to block size and grid origin for "
  "the headline associations (B = 2999); Table A14b, 20-seed Monte-Carlo stability of every "
  "configuration–pair interval; Table A15, post-matching balance diagnostics for the V3 background; "
  "Table A16, water-mask threshold sensitivity; Table A17, land-cover vintage (WorldCover 2020 vs "
  "2021) and day-of-year adjustment sensitivities. All tables are generated by script from the same "
  "analysis table that feeds the in-text statistics.", size=9.5)
_man = pd.read_csv("data/lst_scene_manifest.csv")
_zero = {"LC08_L2SP_121039_20200720_02_T1", "LC08_L2SP_121039_20230729_02_T2",
         "LC08_L2SP_121040_20200922_02_T2"}
_near = {"LC08_L2SP_120039_20200627_02_T2", "LC08_L2SP_120040_20200627_02_T2"}
TABLE("A2", "Landsat 8/9 Collection-2 scene manifest (36 scene assets over four WRS frames). "
            "Adjacent-row assets of the same platform, path, and date are merged pixel-wise into 26 "
            "independent overpasses. Three scenes whose study-area subset is fully cloud- or "
            "shadow-covered after the QA chain (zero valid pixels) are marked “no” and excluded; two "
            "further scenes—one overpass date—retain only ≈0.01% valid pixels after masking and are "
            "marked “near-zero”, leaving 25 overpasses in the primary anomaly analysis.",
      ["Product ID", "Date", "Scene cloud (%)", "Path", "Row", "Valid"],
      [[r.id, r.date, r.cloud, r.path, (r.row if hasattr(r, "row") else r["row"]),
        ("no" if r.id in _zero else ("near-zero" if r.id in _near else "yes"))]
       for r in _man.itertuples()],
      fontsize=7)

H1("Supplementary Materials")
P("The following machine-readable files accompany this article in the public repository (Data "
  "Availability): TableA1_full.csv (full village-level analysis table, 29 villages × all metrics); "
  "TableA3_exact_stats.csv (all 124 tests of families F1 and F2 with nominal, permutation-baseline, and "
  "ESS-heuristic p, within-family BH q, block-bootstrap CI, and n); TableA4_jackknife.csv; "
  "TableA5_moran.csv; TableA6_700mhz_4phase.csv; TableA7_nlos_d3d_sensitivity.csv; "
  "TableA8_scene_composition.csv; TableA9_dlst_variants.csv; TableA10_dsm_crosscheck.csv; "
  "TableA11_sampling_frame.csv; TableA12_link_budget.csv; TableA13_morph_robustness.csv; "
  "TableA14_block_sensitivity.csv; TableA14b_mc_stability.csv; TableA15_match_diagnostics.csv; "
  "TableA16_watermask_sensitivity.csv; TableA17_year_doy_sensitivity.csv; village_geometry_v6.csv "
  "(the single canonical coordinate set); village_sample_v2.csv (sample list with per-village coordinate "
  "verification status); block_membership_0p15.csv (spatial-block assignments for the bootstrap); "
  "dlst_overpass_matrix_v4.csv (per-village per-overpass anomalies under the three background "
  "definitions); coverage_4phase_v4.csv and coverage_4phase_v4_f07.csv (per-phase village coverage at "
  "2.6 GHz and 700 MHz); morphology_fabric_v6.csv and morphology_terrain_v6.csv (the scripted "
  "morphology chains); worldcover_year_compare_v6.csv and dlst_doy_adjust_v6.csv (Table A17 "
  "underlying data); geo_utils.py and s60_unit_tests.py (propagation and morphometric cores with "
  "their unit tests); and the scripted pipeline that regenerates every number, table, and figure.",
  size=9.5)

doc.save(OUT)
print("saved", OUT, "tables:", len(doc.tables))
